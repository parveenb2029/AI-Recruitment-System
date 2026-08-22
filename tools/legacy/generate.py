# =============================================================================
# SUPERSEDED - DO NOT RUN
#
# This script overwrites the ENTIRE project tree from a hardcoded path
# (C:\AI-Recruitment-System). Running it destroys every hand-edited
# document, with no recovery path outside git.
#
# As of the initial commit the documentation is hand-maintained. This file is
# kept only as a record of how the original documentation was produced.
#
# See CLAUDE.md -> "Hard rules".
# =============================================================================

#!/usr/bin/env python3
"""Generate complete AI Recruitment Operations documentation project."""
import json
import os
from datetime import datetime
from pathlib import Path

ROOT = Path(r"C:\AI-Recruitment-System")

WORKFLOWS = {
    "01_Job_Descriptions": {
        "id": "WF-01",
        "name": "Job Description Parsing and Management",
        "short": "JD Parser",
        "purpose": "Standardize, parse, and govern job descriptions before they enter the recruitment pipeline.",
        "trigger": "Hiring manager submits or updates a job requisition in the ATS or shared intake folder.",
        "inputs": ["Raw JD (PDF/DOCX/TXT)", "Requisition metadata", "Department budget code", "Hiring manager profile"],
        "outputs": ["Structured JD JSON", "Validated competency model", "Published JD artifact", "Audit log entry"],
        "actors": ["Hiring Manager", "HR Business Partner", "Talent Acquisition Specialist", "AI JD Parser", "Compensation Analyst"],
        "ai_role": "Extract role title, level, must-have vs nice-to-have skills, experience bands, location, and compensation hints into a canonical schema.",
    },
    "02_Incoming_Resumes": {
        "id": "WF-02",
        "name": "Incoming Resume Intake",
        "short": "Resume Intake",
        "purpose": "Accept, deduplicate, virus-scan, and route candidate applications into the processing queue.",
        "trigger": "Candidate applies via careers portal, email alias, or recruiter upload.",
        "inputs": ["Resume file", "Application form data", "Source channel", "Consent flags", "Referral metadata"],
        "outputs": ["Intake receipt ID", "Quarantined or accepted file", "Candidate master key", "Routing decision"],
        "actors": ["Candidate", "Recruiter", "ATS Webhook", "Email Parser", "Security Scanner", "HR Coordinator"],
        "ai_role": "Classify document type, detect duplicates against existing candidate records, and assign priority based on referral or critical role flags.",
    },
    "03_Extracted_Data": {
        "id": "WF-03",
        "name": "Resume Data Extraction",
        "short": "Resume Extractor",
        "purpose": "Transform unstructured resumes into validated, searchable candidate profiles.",
        "trigger": "Accepted resume lands in processing queue with linked requisition ID.",
        "inputs": ["Resume PDF/DOCX", "OCR fallback flag", "Candidate ID", "Target JD schema version"],
        "outputs": ["Candidate profile JSON", "Confidence scores", "Extraction audit trail", "Human review queue item"],
        "actors": ["AI Resume Parser", "OCR Service", "Data Validator", "Recruiter", "Privacy Officer"],
        "ai_role": "Extract contact info, employment history, education, skills, certifications, and project highlights with field-level confidence.",
    },
    "04_Match_Results": {
        "id": "WF-04",
        "name": "JD-Candidate Matching",
        "short": "Matcher",
        "purpose": "Score and explain fit between structured job requirements and candidate profiles.",
        "trigger": "Validated candidate profile and active JD JSON both available for same requisition.",
        "inputs": ["Structured JD", "Candidate profile JSON", "Matching rubric weights", "DEI guardrails config"],
        "outputs": ["Match score report", "Gap analysis", "Evidence citations", "Ranked candidate list"],
        "actors": ["AI Matcher", "Hiring Manager", "TA Specialist", "DEI Reviewer", "Analytics Engine"],
        "ai_role": "Compute weighted fit scores, surface skill gaps, and generate explainable rationale tied to resume evidence.",
    },
    "05_Shortlisted": {
        "id": "WF-05",
        "name": "Candidate Shortlisting",
        "short": "Shortlist",
        "purpose": "Convert match results into an approved interview slate with documented rationale.",
        "trigger": "Match scores published and SLA window for initial review elapsed or manual review requested.",
        "inputs": ["Match reports", "Recruiter notes", "HM shortlist criteria", "Interview capacity calendar"],
        "outputs": ["Approved shortlist", "Rejected with reason codes", "HM sign-off record", "Interview scheduling tasks"],
        "actors": ["Recruiter", "Hiring Manager", "AI Summarizer", "Interview Scheduler", "HRBP"],
        "ai_role": "Draft shortlist recommendations with risk flags, diversity slate checks, and concise candidate summaries for HM review.",
    },
    "06_Interview_Questions": {
        "id": "WF-06",
        "name": "Interview Question Generation",
        "short": "Interview Gen",
        "purpose": "Produce role-calibrated, legally compliant interview guides for each shortlisted candidate.",
        "trigger": "Candidate approved for interview stage and interview loop defined.",
        "inputs": ["JD JSON", "Candidate profile", "Interview stage (screen/tech/behavioral)", "Competency framework", "Prior interview notes"],
        "outputs": ["Stage-specific question sets", "Scoring rubric", "Interviewer briefing doc", "Anti-bias checklist"],
        "actors": ["AI Interview Generator", "Hiring Manager", "Interviewer Panel", "Legal/Compliance", "Recruiter"],
        "ai_role": "Generate tailored behavioral and technical questions mapped to competencies, with follow-up probes and red-flag guidance.",
    },
    "07_Interview_Feedback": {
        "id": "WF-07",
        "name": "Interview Feedback Processing",
        "short": "Feedback Processor",
        "purpose": "Normalize, synthesize, and validate interviewer feedback into decision-ready artifacts.",
        "trigger": "Interviewer submits feedback form or transcript within 24 hours of interview.",
        "inputs": ["Interviewer scorecards", "Interview transcript (optional)", "Competency rubric", "Calibration guidelines"],
        "outputs": ["Normalized feedback JSON", "Consensus summary", "Hire/no-hire signal", "Calibration flags"],
        "actors": ["Interviewers", "AI Feedback Synthesizer", "Recruiting Coordinator", "Hiring Manager", "Bar Raiser"],
        "ai_role": "Extract ratings, summarize themes, detect conflicting signals, and highlight evidence-backed strengths and concerns.",
    },
    "08_Final_Decision": {
        "id": "WF-08",
        "name": "Final Hiring Decision",
        "short": "Final Decision",
        "purpose": "Consolidate all hiring signals into an auditable offer or rejection recommendation.",
        "trigger": "All interview stages complete and feedback SLA met.",
        "inputs": ["Feedback synthesis", "Match history", "Compensation band", "Background check status", "Headcount approval"],
        "outputs": ["Hiring recommendation report", "Offer letter draft or rejection notice", "Audit package", "Onboarding trigger"],
        "actors": ["Hiring Manager", "HRBP", "Compensation", "Legal", "AI Decision Packager", "Executive Approver"],
        "ai_role": "Compile decision dossier, draft offer/rejection communications, and flag compliance or equity anomalies for human approval.",
    },
}

def write(path: Path, content: str):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")

def workflow_spec(folder: str, wf: dict) -> str:
    return f"""# {wf['name']} — Workflow Specification

**Workflow ID:** {wf['id']}  
**Version:** 1.0.0  
**Last Updated:** {datetime.now().strftime('%Y-%m-%d')}  
**Document Owner:** AI Operations — Talent Technology  
**Classification:** Internal — Operational  

---

## 1. Purpose

{wf['purpose']} This workflow is a core stage in the enterprise AI Recruitment Operations platform. It ensures that every downstream automation receives consistent, auditable inputs while preserving mandatory human oversight at defined approval gates.

The workflow exists because recruitment at scale cannot rely on ad hoc documents, inbox-driven processes, or inconsistent reviewer judgment. By codifying triggers, actors, and success criteria, the organization reduces time-to-fill, improves candidate experience, and maintains defensible hiring records for compliance audits.

---

## 2. Business Goal

| Goal | Target | Measurement Horizon |
|------|--------|---------------------|
| Reduce manual handling time | ≥ 40% vs baseline | Quarterly |
| Improve data consistency | ≥ 95% schema compliance | Monthly |
| Maintain compliance | Zero critical audit findings | Annual |
| Accelerate pipeline velocity | ≥ 25% faster stage transitions | Quarterly |

The business goal is not full automation—it is **augmented operations**: AI handles extraction, matching, and drafting; humans retain authority over shortlist, interview design approval, and final hire decisions.

---

## 3. Business Problem

### 3.1 Current State

Recruiting teams receive unstructured inputs from multiple channels. Hiring managers write job descriptions in inconsistent formats. Recruiters manually copy data into spreadsheets and ATS fields. Interview feedback arrives as free text with incompatible rating scales. Decision meetings lack a single source of truth.

### 3.2 Impact

- **Time:** Recruiters spend 60–70% of time on administrative tasks rather than candidate engagement.
- **Quality:** Inconsistent JD language leads to mismatched applicants and rework.
- **Risk:** Incomplete audit trails expose the organization during EEO or privacy reviews.
- **Experience:** Candidates wait days for acknowledgment because intake is manual.

### 3.3 Root Causes

1. Fragmented tooling (email, SharePoint, ATS, calendars).
2. No canonical data model across hiring stages.
3. Prompt and automation sprawl without governance.
4. Unclear RACI between HR, TA, and hiring managers.

---

## 4. Current Manual Process

1. Hiring manager emails a Word document to recruiter.
2. Recruiter reformats JD and posts to careers site manually.
3. Applications arrive via portal and email; recruiter downloads attachments.
4. Recruiter reads each resume and updates ATS fields by hand.
5. HM receives unstructured resume summaries via email.
6. Interviewers invent questions independently before each loop.
7. Feedback collected in separate forms; recruiter compiles in spreadsheet.
8. Offer drafted from template with manual copy-paste errors.

**Average cycle time (manual baseline):** 18–25 business days from application to offer for mid-level roles.

---

## 5. Pain Points

| Pain Point | Severity | Frequency | Owner |
|------------|----------|-----------|-------|
| Duplicate candidate records | High | Daily | TA Ops |
| Missing mandatory JD fields | High | Weekly | HRBP |
| Low-confidence extractions accepted silently | Critical | Weekly | AI Ops |
| HM bottleneck on shortlist | Medium | Daily | Hiring Manager |
| Inconsistent interview rubrics | High | Per requisition | TA Lead |
| Delayed feedback submission | Medium | 30% of loops | Interviewers |

Additional pain points specific to **{wf['short']}**:

- Inputs arrive without required metadata, blocking automation.
- Reviewers lack explainability for AI-generated recommendations.
- Legacy files (scanned PDFs) break parsing assumptions.

---

## 6. Proposed AI Solution

{wf['ai_role']}

### 6.1 Design Principles

1. **Human-in-the-loop by default** for decisions affecting candidate outcomes.
2. **Schema-first** — all AI outputs validate against versioned JSON schemas.
3. **Explainability** — every score or recommendation cites source evidence.
4. **Fail safe** — low confidence routes to human review, never silent pass-through.
5. **Auditability** — prompt version, model ID, and reviewer ID logged per artifact.

### 6.2 Components

| Component | Responsibility |
|-----------|----------------|
| Intake Gateway | Accept files, assign IDs, enforce virus scan |
| AI Orchestrator | Invoke prompts, enforce timeouts, retry policy |
| Validation Service | Schema + business rule checks |
| Review Console | HR/HM approval UI |
| Artifact Store | Immutable storage under `{folder}/` |
| Metrics Emitter | KPI events to analytics warehouse |

---

## 7. Workflow Trigger

**Primary trigger:** {wf['trigger']}

**Secondary triggers:**
- Scheduled batch reprocessing (nightly) for failed items in dead-letter queue.
- Manual re-run by TA Specialist with `force=true` and documented reason.
- Upstream workflow completion event (webhook `workflow.completed`).

**Preconditions:**
- Active requisition exists with valid department and headcount approval.
- Candidate consent on file where required by jurisdiction.
- Prompt library version compatible with orchestrator runtime.

---

## 8. Inputs

{chr(10).join(f'- {i}' for i in wf['inputs'])}

### 8.1 Input Quality Requirements

| Input | Format | Max Size | Required Fields |
|-------|--------|----------|-----------------|
| Primary document | PDF, DOCX, TXT | 10 MB | Readable text layer or OCR path |
| Metadata envelope | JSON | 50 KB | `requisition_id`, `source`, `timestamp` |
| Actor context | JSON | 10 KB | `submitter_email`, `role` |

---

## 9. Outputs

{chr(10).join(f'- {o}' for o in wf['outputs'])}

Outputs are written to `{folder}/` with naming convention:  
`{{requisition_id}}_{{candidate_id}}_{{artifact_type}}_{{iso_timestamp}}.{{ext}}`

---

## 10. Actors

{chr(10).join(f'- **{a}**' for a in wf['actors'])}

### RACI Matrix (Summary)

| Activity | TA Specialist | Hiring Manager | AI System | HRBP | Legal |
|----------|---------------|----------------|-----------|------|-------|
| Submit/trigger | R/A | C | I | I | I |
| AI processing | I | I | R/A | I | I |
| Quality review | R | C | I | A | I |
| Approve stage output | C | A | I | R | C |
| Exception override | R | A | I | C | C |

*R = Responsible, A = Accountable, C = Consulted, I = Informed*

---

## 11. Dependencies

| Dependency | Type | Failure Impact |
|------------|------|----------------|
| Upstream workflow completion | Hard | Blocks start |
| Prompt Library v1.x | Hard | Blocks AI step |
| JSON Schema Registry | Hard | Blocks validation |
| ATS API (Workday/Greenhouse) | Soft | Manual fallback |
| Identity Provider (SSO) | Hard | Blocks human review |
| Object storage (SharePoint/S3) | Hard | Blocks artifact persistence |

**Upstream:** Prior workflow in pipeline must emit valid handoff envelope.  
**Downstream:** Next workflow consumes validated outputs only.

---

## 12. Step-by-Step Workflow

### Phase A — Initiation
1. Trigger fires; orchestrator creates `workflow_run_id`.
2. Load workflow config and active prompt versions from registry.
3. Validate inputs (file type, size, required metadata).
4. Write intake record to audit log with SHA-256 of source file.

### Phase B — AI Processing
5. Pre-process document (text extraction or OCR if needed).
6. Invoke AI with system + user prompts from `Prompt.md`.
7. Parse response; attempt JSON extraction if structured output expected.
8. Run schema validation and business rule engine.
9. Compute confidence aggregate; compare to threshold (default 0.85).

### Phase C — Human Review (Conditional)
10. If confidence < threshold OR mandatory review flag → enqueue review task.
11. Reviewer accepts, edits, or rejects with reason code.
12. Edited outputs re-validate before release.

### Phase D — Publication
13. Persist approved artifacts to `{folder}/`.
14. Emit `workflow.completed` event to downstream subscribers.
15. Update ATS status via API or manual task if API unavailable.
16. Record KPI metrics (latency, automation rate, review rate).

---

## 13. Business Rules

1. **BR-01:** No candidate-facing action without HR review for rejection templates.
2. **BR-02:** Duplicate candidates merged under single `candidate_master_id`.
3. **BR-03:** JD must include EEO statement before external publish.
4. **BR-04:** Match scores below 0.40 auto-archive unless recruiter overrides with justification.
5. **BR-05:** All AI outputs retain prompt version and model ID for 7 years.
6. **BR-06:** PII fields masked in logs; full data only in secured artifact store.
7. **BR-07:** {wf['short']}-specific: processing SLA = 4 business hours for standard priority.

---

## 14. Decision Points

| Decision ID | Question | Default Path | Escalation |
|-------------|----------|--------------|------------|
| D-01 | Input valid? | Yes → continue | No → reject + notify submitter |
| D-02 | AI confidence ≥ threshold? | Yes → auto-publish | No → human review |
| D-03 | Schema valid? | Yes → continue | No → retry once, then DLQ |
| D-04 | Duplicate detected? | Merge | Flag for recruiter |
| D-05 | Compliance flag raised? | Hold | Legal review within 24h |

---

## 15. Approval Points

| Gate | Approver | SLA | Outcome |
|------|----------|-----|---------|
| AP-01 | TA Specialist | 4h | Accept AI output for non-decision artifacts |
| AP-02 | Hiring Manager | 24h | Shortlist and interview plan approval |
| AP-03 | HRBP | 48h | Offer/comp band approval |
| AP-04 | Legal | 72h | Non-standard offer terms |

**Never fully automated:** Final hire/reject decision, compensation exceptions, adverse impact overrides.

---

## 16. Exception Handling

| Exception | Detection | Response | Recovery |
|-----------|-----------|----------|----------|
| Invalid file | Magic byte check | Reject + email template | Candidate re-upload |
| OCR failure | Empty text layer | Route to manual transcription | Recruiter entry |
| AI timeout | 120s limit | Exponential backoff retry ×3 | DLQ + alert |
| Invalid JSON | Schema validator | Repair prompt or human fix | Versioned re-run |
| ATS sync failure | HTTP 5xx | Queue retry job | Manual ATS update task |

All exceptions logged with correlation ID linking to `workflow_run_id`.

---

## 17. Success Criteria

1. ≥ 90% of standard-priority items complete within SLA without manual intervention.
2. ≥ 95% of published artifacts pass schema validation on first attempt after review.
3. Zero PII leakage incidents attributable to this workflow.
4. Hiring manager satisfaction ≥ 4.0/5.0 on output usefulness survey.
5. Audit simulation: 100% of sample records traceable to source file and prompt version.

---

## 18. KPIs

| KPI | Definition | Target | Data Source |
|-----|------------|--------|-------------|
| Processing time (P50) | Trigger to published artifact | < 15 min | Orchestrator |
| Processing time (P95) | Same | < 2 hours | Orchestrator |
| Automation rate | Completed without human edit | ≥ 70% | Review console |
| First-pass validation rate | Schema valid post-AI | ≥ 85% | Validator |
| Human review rate | Items entering review queue | ≤ 30% | Review console |
| Error rate | Failed runs / total runs | < 2% | DLQ metrics |
| Rework rate | Re-runs after rejection | < 5% | Audit log |

---

## 19. Risks

| Risk | Likelihood | Impact | Mitigation |
|------|------------|--------|------------|
| Model hallucination on skills | Medium | High | Evidence citations + human review |
| Bias in matching | Medium | Critical | DEI guardrails, regular bias audits |
| Data privacy breach | Low | Critical | Encryption, RBAC, retention policy |
| Over-automation perception | Medium | Medium | Clear RACI, reviewer training |
| Vendor API change | Medium | Medium | Abstraction layer, contract SLAs |
| Prompt drift after model update | High | Medium | Prompt testing SOP before rollout |

---

## 20. Future Improvements

1. **Multimodal parsing** — portfolio links, GitHub, LinkedIn enrichment with consent.
2. **Active learning** — reviewer corrections feed fine-tuning dataset with governance.
3. **Real-time collaboration** — HM and recruiter co-edit shortlist in shared console.
4. **Predictive analytics** — forecast time-to-fill based on funnel metrics.
5. **Cross-lingual support** — parse resumes in Hindi, Spanish, French with locale-specific schemas.
6. **Integration expansion** — native Workday Recruiting bi-directional sync.

---

## Appendix A — Glossary

| Term | Definition |
|------|------------|
| DLQ | Dead-letter queue for failed automation runs |
| HM | Hiring Manager |
| TA | Talent Acquisition |
| JD | Job Description |
| SLA | Service Level Agreement |

## Appendix B — Related Documents

- `{folder}/Prompt.md`
- `{folder}/Execution_Flow.md`
- `{folder}/Automation.md`
- `09_Prompt_Library/` — reusable prompts
- `10_SOPs/` — operational procedures
- `schemas/` — JSON schema definitions

---

*End of Workflow Specification — {wf['id']}*
"""


def prompt_md(folder: str, wf: dict) -> str:
    prompt_name = wf["short"].upper().replace(" ", "_")
    return f"""# {wf['name']} — AI Prompt Specification

**Prompt ID:** PROMPT-{wf['id']}  
**Prompt Version:** 1.0.0  
**Prompt Owner:** AI Operations — Prompt Engineering Lead  
**Last Reviewed:** {datetime.now().strftime('%Y-%m-%d')}  
**Model Compatibility:** GPT-4.1+, Claude 3.5+, Azure OpenAI gpt-4o  

---

## Document Purpose

This document defines the production prompt used in **{wf['name']}**. Every section exists to ensure repeatability, auditability, and safe deployment across environments. Prompts are treated as **infrastructure-as-code**: versioned, tested, and owned.

---

## Why Each Section Exists

| Section | Why It Exists |
|---------|---------------|
| System Prompt | Sets immutable behavioral guardrails, tone, compliance constraints, and output format the model must never violate. |
| User Prompt | Carries task-specific variables and the actual document/data to process for this workflow run. |
| Expected Output | Defines the contract downstream validators and humans rely on; prevents ambiguous free-text when structure is required. |
| JSON Schema | Enables automated validation, ATS mapping, and type-safe integrations. |
| Validation Rules | Business logic beyond JSON Schema (cross-field checks, confidence thresholds). |
| Examples | Few-shot anchors reduce format drift and establish quality bar for evals. |
| Edge Cases | Documents known failure modes so operators and evaluators test them explicitly. |
| Prompt Version | Supports rollback, A/B testing, and audit ("which prompt produced this artifact?"). |
| Prompt Owner | Accountability for changes, incidents, and approval workflow per Prompt Governance SOP. |
| Prompt Notes | Operational context: known limitations, tuning history, deployment caveats. |

---

## System Prompt

```
You are an enterprise recruitment AI assistant specialized in {wf['short'].lower()} for a global technology employer.

ROLE AND BOUNDARIES:
- You process hiring-related documents and data with precision, neutrality, and compliance awareness.
- You NEVER invent credentials, employment dates, scores, or interview outcomes not supported by source material.
- You NEVER infer protected characteristics (age, race, gender, religion, disability, national origin) or use them in recommendations.
- You flag uncertainty explicitly using confidence scores between 0.0 and 1.0 per extracted field or judgment.

OUTPUT REQUIREMENTS:
- Respond ONLY with valid JSON matching the schema provided in the user message unless explicitly asked for prose.
- Include an "evidence" array citing verbatim snippets (max 25 words each) supporting key conclusions.
- Include "prompt_version": "1.0.0" and "workflow_id": "{wf['id']}" in every response.
- Use ISO 8601 dates. Use canonical skill names from the provided skills ontology when available.

COMPLIANCE:
- Apply GDPR/CCPA minimization: do not repeat full government IDs, full DOB, or unnecessary PII in summaries.
- If source document is unreadable or empty, return status "FAILED" with reason code — do not guess.

TONE:
- Professional, concise, audit-ready language suitable for HR systems and hiring manager review.
```

**Why:** The system prompt is the security and compliance kernel. It applies to every invocation regardless of input variation.

---

## User Prompt Template

```
Process the following input for workflow {wf['id']} ({wf['name']}).

=== METADATA ===
requisition_id: {{{{requisition_id}}}}
candidate_id: {{{{candidate_id}}}}
workflow_run_id: {{{{workflow_run_id}}}}
source_channel: {{{{source_channel}}}}
priority: {{{{priority}}}}

=== PRIMARY CONTENT ===
{{{{document_text_or_structured_input}}}}

=== REFERENCE SCHEMA VERSION ===
{{{{schema_version}}}}

=== OPTIONAL CONTEXT ===
{{{{additional_context}}}}

Return JSON conforming to the {prompt_name}_OUTPUT schema.
Include field-level confidence scores for all extracted or inferred fields.
```

**Why:** Separating metadata from content allows the orchestrator to inject variables safely and redact PII in logs while keeping full content in the model call.

---

## Expected Output

```json
{{
  "status": "SUCCESS | PARTIAL | FAILED",
  "workflow_id": "{wf['id']}",
  "prompt_version": "1.0.0",
  "requisition_id": "REQ-2026-0142",
  "candidate_id": "CAN-88421",
  "processed_at": "2026-08-03T10:30:00Z",
  "confidence_aggregate": 0.92,
  "results": {{}},
  "evidence": [
    {{"field": "example_field", "snippet": "verbatim source text", "source_location": "page 1"}}
  ],
  "flags": [],
  "human_review_required": false,
  "review_reasons": []
}}
```

**Why:** Downstream automation branches on `status`, `human_review_required`, and `confidence_aggregate`. Partial success is first-class to avoid silent data loss.

---

## JSON Schema

See `schemas/` for canonical definitions. Workflow-specific output extends base schemas:

```json
{{
  "$schema": "https://json-schema.org/draft/2020-12/schema",
  "$id": "https://recruitment.example.com/schemas/{wf['id']}_output.json",
  "type": "object",
  "required": ["status", "workflow_id", "prompt_version", "confidence_aggregate"],
  "properties": {{
    "status": {{"enum": ["SUCCESS", "PARTIAL", "FAILED"]}},
    "workflow_id": {{"const": "{wf['id']}"}},
    "prompt_version": {{"type": "string"}},
    "confidence_aggregate": {{"type": "number", "minimum": 0, "maximum": 1}},
    "human_review_required": {{"type": "boolean"}},
    "results": {{"type": "object"}},
    "evidence": {{
      "type": "array",
      "items": {{
        "type": "object",
        "required": ["field", "snippet"],
        "properties": {{
          "field": {{"type": "string"}},
          "snippet": {{"type": "string", "maxLength": 200}}
        }}
      }}
    }}
  }}
}}
```

**Why:** Schema validation catches malformed model output before it reaches ATS or candidate-facing channels.

---

## Validation Rules

| Rule ID | Rule | Action on Fail |
|---------|------|----------------|
| VR-01 | `confidence_aggregate` ≥ 0.85 for auto-publish | Route to human review |
| VR-02 | All required schema fields present | Retry with repair prompt once |
| VR-03 | Evidence snippets must appear in source text (fuzzy match ≥ 0.8) | Flag `POTENTIAL_HALLUCINATION` |
| VR-04 | Dates must be chronologically consistent | Flag for reviewer |
| VR-05 | No protected-class terms in recommendation fields | Block + compliance alert |
| VR-06 | `prompt_version` matches deployed version | Reject response |

---

## Examples

### Example 1 — Successful Processing

**Input excerpt:** Standard software engineer resume with clear employment history.  
**Output excerpt:**
```json
{{"status": "SUCCESS", "confidence_aggregate": 0.94, "human_review_required": false}}
```

### Example 2 — Partial Success (OCR noise)

**Input excerpt:** Scanned PDF with garbled dates in one role.  
**Output excerpt:**
```json
{{"status": "PARTIAL", "confidence_aggregate": 0.71, "human_review_required": true, "review_reasons": ["LOW_CONFIDENCE_DATES"]}}
```

### Example 3 — Failed (Empty document)

**Output excerpt:**
```json
{{"status": "FAILED", "confidence_aggregate": 0.0, "human_review_required": true, "review_reasons": ["EMPTY_SOURCE"]}}
```

**Why:** Examples anchor evaluation suites and onboarding for new prompt engineers.

---

## Edge Cases

| Edge Case | Expected Behavior |
|-----------|-------------------|
| Password-protected PDF | Return FAILED, reason `ENCRYPTED_DOCUMENT` |
| Non-English resume | Process if readable; set `locale_detected`; lower confidence if untranslated |
| 20+ page CV | Summarize oldest roles; retain last 10 years in detail |
| Candidate applies to multiple reqs | Include `requisition_id` in output; never merge profiles automatically |
| Conflicting dates | List conflict in `flags`; never silently resolve |
| Model returns markdown fences | Strip fences in post-processor before JSON parse |

---

## Prompt Version History

| Version | Date | Author | Changes |
|---------|------|--------|---------|
| 1.0.0 | 2026-08-03 | AI Ops Team | Initial production release |

---

## Prompt Owner

**Primary:** Prompt Engineering Lead — Talent Technology  
**Backup:** Senior AI Operations Specialist  
**Approval required from:** HR Compliance (material changes affecting candidate evaluation)

---

## Prompt Notes

- Temperature recommendation: **0.1** for extraction tasks; **0.3** for question generation.
- Max tokens: 4096 output; increase to 8192 for senior executive resumes.
- Known limitation: tables in multi-column PDFs may misalign skills — prefer DOCX intake when possible.
- Run regression eval (`Prompt Testing SOP`) before any version bump.
- Do not deploy prompt changes during active offer approval windows without change advisory.

---

*End of Prompt Specification — PROMPT-{wf['id']}*
"""


def execution_flow_md(folder: str, wf: dict) -> str:
    return f"""# {wf['name']} — Execution Flow

**Workflow ID:** {wf['id']}  
**Version:** 1.0.0  

---

## 1. Overview

This document describes the runtime execution path for **{wf['name']}**, including automation entry points, AI invocation boundaries, human review gates, storage locations, and failure handling.

---

## 2. Trigger Sources

| Source | Mechanism | Priority |
|--------|-----------|----------|
| ATS webhook | REST POST `/hooks/{folder}` | Standard |
| Email intake | Parser → object storage | Standard |
| Manual upload | SharePoint drop folder | High |
| Upstream workflow | Event bus `workflow.completed` | Standard |
| Scheduled retry | Cron every 15 min (DLQ) | Low |

---

## 3. Primary Execution Flow

```mermaid
flowchart TD
    A[Trigger Received] --> B{{Input Validation}}
    B -->|Invalid| C[Reject + Notify Submitter]
    B -->|Valid| D[Create workflow_run_id]
    D --> E[Load Prompt v1.0.0]
    E --> F[Extract Text / OCR]
    F --> G[AI Processing]
    G --> H{{JSON Valid?}}
    H -->|No| I[Retry / Repair Prompt]
    I --> G
    H -->|Yes| J{{Confidence >= 0.85?}}
    J -->|Yes| K[Auto-Publish Artifact]
    J -->|No| L[Human Review Queue]
    L --> M{{Approved?}}
    M -->|Yes| K
    M -->|No| N[Reject + Reason Code]
    K --> O[Write to {folder}/]
    O --> P[Emit workflow.completed]
    P --> Q[Update ATS Status]
    C --> R[Log + Metrics]
    N --> R
    Q --> R
```

---

## 4. Sequence Diagram — AI Invocation

```mermaid
sequenceDiagram
    participant T as Trigger
    participant O as Orchestrator
    participant S as Storage
    participant AI as LLM API
    participant V as Validator
    participant H as Human Reviewer
    participant ATS as ATS

    T->>O: New work item
    O->>S: Fetch source document
    O->>AI: System + User prompt
    AI-->>O: JSON response
    O->>V: Schema + business rules
    alt Validation pass, high confidence
        O->>S: Save artifact
        O->>ATS: Status update
    else Low confidence
        O->>H: Review task
        H-->>O: Approve / Edit / Reject
        O->>S: Save approved artifact
    end
```

---

## 5. Storage Map

| Stage | Location | Retention |
|-------|----------|-----------|
| Raw intake | `02_Incoming_Resumes/` (if applicable) | 7 years |
| Processed output | `{folder}/` | 7 years |
| Failed items | `DLQ/{folder}/` | 90 days |
| Audit logs | SIEM / Log Analytics | 7 years |
| Review decisions | Review Console DB | 7 years |

---

## 6. Failure Paths

```mermaid
flowchart LR
    F1[AI Timeout] --> R1[Retry x3]
    R1 --> DLQ[Dead Letter Queue]
    F2[Invalid JSON] --> R2[Repair Prompt]
    R2 --> DLQ
    F3[OCR Failure] --> HR[Manual Transcription Task]
    F4[Duplicate] --> MERGE[Merge Workflow]
    DLQ --> OPS[On-call Alert]
```

---

## 7. Latency Budget

| Step | Target P50 | Target P95 |
|------|------------|------------|
| Validation | 2s | 5s |
| Text extraction | 5s | 30s |
| AI call | 15s | 90s |
| Schema validation | 1s | 3s |
| Human review | 4h SLA | 24h SLA |
| **Total automated path** | **< 2 min** | **< 5 min** |

---

## 8. Idempotency

- All runs keyed by `workflow_run_id` + `content_hash`.
- Re-trigger with same hash returns existing artifact unless `force=true`.
- Prevents duplicate ATS updates and double-charging API usage.

---

*End of Execution Flow — {wf['id']}*
"""


def architecture_md(folder: str, wf: dict) -> str:
    return f"""# {wf['name']} — Architecture

**Workflow ID:** {wf['id']}  
**Version:** 1.0.0  

---

## 1. Architectural Context

This workflow operates as a **stage** within the modular recruitment pipeline. It adheres to event-driven, schema-first design with clear separation between ingestion, intelligence, validation, and human approval layers.

---

## 2. Component Diagram

```mermaid
flowchart TB
    subgraph Presentation
        SP[SharePoint Drop Zone]
        RC[Review Console]
        EM[Email Intake]
    end
    subgraph Orchestration
        OR[Workflow Orchestrator]
        EB[Event Bus]
        DLQ[Dead Letter Queue]
    end
    subgraph Intelligence
        LLM[LLM Gateway]
        OCR[OCR Service]
        PR[Prompt Registry]
    end
    subgraph Data
        AS[Artifact Store]
        SCH[Schema Registry]
        ATS[Applicant Tracking System]
    end
    SP --> OR
    EM --> OR
    OR --> OCR
    OR --> LLM
    PR --> LLM
    LLM --> OR
    OR --> SCH
    OR --> AS
    OR --> RC
    RC --> OR
    OR --> ATS
    OR --> EB
    OR --> DLQ
```

---

## 3. Layer Responsibilities

| Layer | Components | Responsibility |
|-------|------------|----------------|
| Presentation | SharePoint, Email, Review UI | Human and system entry points |
| Orchestration | n8n / Python / Azure Functions | State machine, retries, routing |
| Intelligence | LLM, OCR, Prompt Registry | Unstructured → structured transformation |
| Data | S3/SharePoint, PostgreSQL, ATS | Persistence, canonical records |
| Governance | Schema Registry, Audit Log | Compliance, traceability |

---

## 4. Technology Choices

| Decision | Choice | Rationale |
|----------|--------|-----------|
| LLM access | Azure OpenAI private endpoint | Enterprise data residency, SLA |
| Orchestration | n8n self-hosted + Python workers | Visual ops + complex logic |
| Storage | SharePoint + blob archive | HR familiarity + immutability |
| Validation | JSON Schema + custom rules engine | Fail fast before ATS write |
| Identity | Azure AD SSO | RBAC aligned to HR org |

---

## 5. Security Architecture

- **Encryption at rest:** AES-256 for all artifact stores.
- **Encryption in transit:** TLS 1.2+ for all API calls.
- **RBAC:** TA Specialist, HM, HRBP, Admin roles with least privilege.
- **PII handling:** Tokenization in logs; full PII only in secured artifact paths.
- **Network:** LLM calls via private endpoint; no public internet for candidate data.

---

## 6. Scalability

- Horizontal scale of stateless Python workers behind queue.
- LLM rate limiting with token bucket per workflow priority.
- Batch mode for nightly reprocessing of DLQ items.

---

## 7. Integration Points

| System | Direction | Protocol |
|--------|-----------|----------|
| Workday/Greenhouse ATS | Bi-directional | REST API |
| SharePoint | Inbound files | Graph API |
| Email (O365) | Inbound | Graph webhook |
| SIEM | Outbound logs | Syslog / HTTP |
| Power BI | Outbound metrics | SQL / API |

---

## 8. Deployment Model

```
Dev → Test (synthetic data only) → Staging (masked prod) → Production
```

Prompt and schema changes require promotion through all environments with eval gates.

---

*End of Architecture — {wf['id']}*
"""


def automation_md(folder: str, wf: dict) -> str:
    return f"""# {wf['name']} — Automation Implementation Guide

**Workflow ID:** {wf['id']}  
**Version:** 1.0.0  

---

## 1. Automation Strategy

This workflow is designed for **composable automation**: each platform handles what it does best. Cursor accelerates prompt iteration; Python handles validation and ATS integration; n8n orchestrates visual workflows; Zapier/Make connect SaaS edges; REST/webhooks tie the ATS and storage layers together.

---

## 2. Component Responsibility Matrix

| Component | Responsibility in {wf['short']} |
|-----------|--------------------------------|
| **Cursor** | Prompt development, eval script authoring, documentation maintenance |
| **Python** | OCR pipeline, JSON validation, ATS SDK, batch reprocessing, metrics |
| **n8n** | Primary orchestrator: triggers, branching, retries, human task creation |
| **Zapier** | Lightweight connectors for teams without n8n (email → Drive) |
| **Make.com** | Complex multi-app scenarios (SharePoint + OpenAI + Slack) |
| **REST APIs** | LLM gateway, ATS, schema registry, review console |
| **Webhooks** | ATS application events, upstream workflow completion |
| **Google Drive** | Alternative intake folder for regional teams |
| **SharePoint** | Primary document library for `{folder}/` |
| **Email** | Application alias parsing, failure notifications |
| **ATS** | System of record for requisition and candidate status |

---

## 3. Cursor Implementation

**Use case:** Prompt engineering and regression testing.

```
1. Open 09_Prompt_Library/ and {folder}/Prompt.md in Cursor.
2. Use Agent to run eval against samples/Rahul_Sharma.json.
3. Commit prompt version bump with eval results in docs/.
```

**Responsibility:** Developer productivity, not production runtime.

---

## 4. Python Implementation

```python
# Pseudocode — production worker for {wf['id']}
def process_item(workflow_run_id, payload):
    doc = fetch_from_sharepoint(payload["path"])
    text = extract_text(doc) or ocr_fallback(doc)
    result = llm_client.complete(system=prompt_system, user=build_user(text, payload))
    validated = validate_schema(result, schema="{wf['id']}_output")
    if validated.confidence < 0.85:
        create_review_task(workflow_run_id, validated)
    else:
        save_artifact(f"{folder}/", validated)
        emit_event("workflow.completed", workflow_run_id)
```

**Libraries:** `pypdf`, `python-docx`, `jsonschema`, `httpx`, `azure-identity`  
**Responsibility:** Heavy logic, idempotency, secure credential handling.

---

## 5. n8n Workflow

```mermaid
flowchart LR
    WH[Webhook Trigger] --> DL[Download File]
    DL --> EX[Execute Python OCR]
    EX --> AI[HTTP Request - LLM]
    AI --> VAL[Function Node - Validate]
    VAL --> IF{{Confidence OK?}}
    IF -->|Yes| SP[SharePoint Upload]
    IF -->|No| TASK[Create Review Task]
    SP --> ATS[ATS Update]
    VAL -->|Error| RETRY[Retry Node]
    RETRY --> DLQ[DLQ Folder]
```

**Responsibility:** Visual orchestration, retry policies, non-developer maintainability.

---

## 6. Zapier / Make.com

| Platform | Pattern | When to Use |
|----------|---------|-------------|
| Zapier | Gmail attachment → Google Drive → Webhook to n8n | Small teams, quick MVP |
| Make.com | SharePoint watch → OpenAI → Teams notification | Microsoft-centric orgs |

**Responsibility:** Edge integration only — not core validation logic.

---

## 7. REST API Contracts

### Inbound Webhook
```
POST /api/v1/workflows/{folder}/run
Authorization: Bearer {{service_token}}
Content-Type: application/json

{{
  "requisition_id": "REQ-2026-0142",
  "candidate_id": "CAN-88421",
  "source_path": "sharepoint://...",
  "priority": "standard"
}}
```

### Outbound Event
```
POST {{subscriber_url}}
{{
  "event": "workflow.completed",
  "workflow_id": "{wf['id']}",
  "workflow_run_id": "RUN-uuid",
  "artifact_path": "{folder}/REQ-2026-0142_CAN-88421_output.json"
}}
```

---

## 8. SharePoint Integration

- **Library:** `RecruitmentOps/{folder}/`
- **Metadata columns:** RequisitionID, CandidateID, ConfidenceScore, PromptVersion
- **Flow:** n8n uploads approved JSON + links source PDF
- **Permissions:** HR-TA-Contributors write; HM read on assigned reqs only

---

## 9. Email Automation

| Event | Template | Recipient |
|-------|----------|-----------|
| Intake rejected | `INTAKE_REJECT` | Candidate / Recruiter |
| Review needed | `REVIEW_REQUIRED` | TA Specialist |
| Stage complete | `STAGE_COMPLETE` | Hiring Manager |

Parse inbound applications via dedicated alias `careers+{folder}@company.com`.

---

## 10. ATS Integration

| Action | ATS Endpoint | Timing |
|--------|--------------|--------|
| Update stage | PATCH /candidates/{{id}}/stage | After publish |
| Attach artifact | POST /candidates/{{id}}/attachments | After publish |
| Log note | POST /candidates/{{id}}/notes | On human review |

Fallback: create recruiter task if API unavailable > 15 min.

---

## 11. Monitoring & Alerting

- n8n execution failures → PagerDuty (P1 if DLQ > 50/hour)
- Python worker health → Kubernetes liveness probe
- LLM latency → Datadog dashboard, alert P95 > 120s

---

*End of Automation Guide — {wf['id']}*
"""


def error_handling_md(folder: str, wf: dict) -> str:
    return f"""# {wf['name']} — Error Handling

**Workflow ID:** {wf['id']}  
**Version:** 1.0.0  

---

## 1. Error Handling Philosophy

**Detect early, fail visibly, recover safely.** No silent degradation of candidate data. Every error produces a correlation ID, structured reason code, and prescribed recovery path.

---

## 2. Error Catalog

| Error | Detection | User Impact | System Response | Recovery |
|-------|-----------|-------------|-----------------|----------|
| **Invalid PDF** | Magic bytes / parse exception | Upload rejected | Notify submitter with format guide | Re-upload as PDF/A or DOCX |
| **Corrupt document** | CRC / truncated file | Processing halted | Quarantine file | Request new copy |
| **Missing fields** | Schema validator | Partial record | `PARTIAL` status + review | Human completes fields |
| **OCR failure** | Empty text, low char count | Delay | Route to manual transcription | Recruiter data entry UI |
| **AI timeout** | 120s no response | Delay | Retry ×3 exponential backoff | DLQ + on-call if persistent |
| **Invalid JSON** | JSON parse error | Delay | Repair prompt attempt once | Human fix or re-run |
| **Duplicate candidate** | Email/hash match | None if merged | Link to existing `candidate_master_id` | Recruiter confirms merge |
| **Unsupported format** | Extension/MIME check | Upload rejected | List supported formats | Convert and re-upload |
| **Network failure** | HTTP timeout/5xx | Delay | Circuit breaker + retry | Queue for batch retry |
| **Human override** | Review console action | Decision changed | Log override reason + actor | Supersedes AI output |

---

## 3. Retry Strategy

```mermaid
flowchart TD
    E[Error Detected] --> C{{Retryable?}}
    C -->|No| DLQ[Dead Letter Queue]
    C -->|Yes| R{{Retry count < 3?}}
    R -->|Yes| W[Wait 2^n seconds]
    W --> X[Retry Operation]
    X --> E
    R -->|No| DLQ
    DLQ --> A[Alert Ops]
    DLQ --> M[Manual Playbook]
```

| Error Type | Max Retries | Backoff |
|------------|-------------|---------|
| AI timeout | 3 | 2s, 4s, 8s |
| Network 5xx | 5 | 1s, 2s, 4s, 8s, 16s |
| Rate limit 429 | 3 | Respect Retry-After header |
| Invalid JSON | 1 | Repair prompt |
| OCR failure | 0 | Immediate human route |

---

## 4. Logging Strategy

### Structured Log Fields
```json
{{
  "timestamp": "ISO8601",
  "level": "ERROR",
  "workflow_id": "{wf['id']}",
  "workflow_run_id": "uuid",
  "correlation_id": "uuid",
  "error_code": "AI_TIMEOUT",
  "message": "LLM did not respond within 120s",
  "retry_count": 2,
  "candidate_id_hash": "sha256_prefix"
}}
```

### Log Destinations
| Destination | Content | Retention |
|-------------|---------|-----------|
| Application logs | All errors | 90 days |
| SIEM | P1/P2 security and compliance | 7 years |
| DLQ metadata | Failed payloads (redacted) | 90 days |
| Review audit | Human overrides | 7 years |

### PII Rules
- Never log full resume text in error logs.
- Log file hash and path only.
- Candidate email masked: `r***@domain.com`

---

## 5. Human Override Protocol

1. Reviewer opens item in Review Console with full context.
2. Selects **Override AI Output** with mandatory reason code.
3. System logs: original AI output, override fields, reviewer ID, timestamp.
4. Override outputs marked `source: HUMAN_OVERRIDE` in artifact metadata.
5. Monthly QA sample of overrides per `AI Quality Review SOP`.

---

## 6. Escalation Matrix

| Severity | Condition | Response Time | Escalate To |
|----------|-----------|---------------|-------------|
| P1 | DLQ flood, PII leak suspect | 15 min | AI Ops on-call + CISO |
| P2 | Workflow down > 1 hour | 30 min | TA Technology Lead |
| P3 | Single item stuck | 4 hours | TA Specialist |
| P4 | Format rejection | Next business day | Recruiter |

---

*End of Error Handling — {wf['id']}*
"""


def validation_md(folder: str, wf: dict) -> str:
    return f"""# {wf['name']} — Validation Framework

**Workflow ID:** {wf['id']}  
**Version:** 1.0.0  

---

## 1. Validation Layers

Validation is applied in **four layers** before any artifact is published or synced to the ATS.

```mermaid
flowchart LR
    F[File Validation] --> D[Data Validation]
    D --> S[Schema Validation]
    S --> B[Business Validation]
    B --> AI[AI Confidence Gate]
    AI --> H[Human Verification]
```

---

## 2. File Validation

| Check | Rule | Fail Action |
|-------|------|-------------|
| Extension | `.pdf`, `.docx`, `.doc`, `.txt` | Reject |
| MIME type | Matches extension magic bytes | Reject |
| Size | ≤ 10 MB | Reject |
| Virus scan | Clean per Defender/API | Quarantine |
| Encryption | Not password-protected | Reject with instructions |
| Readability | ≥ 100 chars extractable OR OCR path | OCR fallback |

---

## 3. Data Validation

| Field | Rule |
|-------|------|
| Email | RFC 5322 format |
| Phone | E.164 or regional format with country |
| Dates | ISO 8601, start ≤ end |
| URLs | HTTPS preferred, malware scan |
| Skills | Normalized against ontology |
| Names | Non-empty, no placeholder values ("Test", "N/A") |

---

## 4. Schema Validation

- Engine: `jsonschema` draft 2020-12
- Schemas loaded from `schemas/` at runtime
- Schema version pinned in workflow config
- Breaking schema changes require migration script

```python
from jsonschema import validate, ValidationError
validate(instance=output, schema=load_schema("{wf['id']}_output"))
```

---

## 5. Business Validation

| Rule | Description |
|------|-------------|
| BV-01 | Requisition must be `OPEN` status |
| BV-02 | Candidate consent flag true for EU/UK |
| BV-03 | No duplicate active application same req |
| BV-04 | Match score components sum to weight 1.0 |
| BV-05 | Interview questions exclude protected-class probes |
| BV-06 | Offer compensation within approved band |

---

## 6. AI Confidence Score

| Aggregate Score | Action |
|-----------------|--------|
| ≥ 0.85 | Auto-publish eligible |
| 0.70 – 0.84 | Mandatory TA review |
| 0.50 – 0.69 | Mandatory TA + spot check |
| < 0.50 | Reject automation; manual processing |

Field-level scores below 0.60 highlighted in review UI.

---

## 7. Human Verification

**Always required for:**
- Final hire/reject decisions
- Compensation outside band
- Any compliance flag
- First 50 runs after prompt version change (per Prompt Testing SOP)

**Spot check sampling:** 10% of auto-published items weekly per AI Quality Review SOP.

---

*End of Validation Framework — {wf['id']}*
"""


def metrics_md(folder: str, wf: dict) -> str:
    return f"""# {wf['name']} — Metrics & KPIs

**Workflow ID:** {wf['id']}  
**Version:** 1.0.0  

---

## 1. Metrics Framework

Metrics align to **efficiency**, **quality**, **compliance**, and **experience** dimensions. All metrics emitted as structured events to the analytics warehouse with dimensions: `workflow_id`, `requisition_id`, `prompt_version`, `region`, `priority`.

---

## 2. Primary KPIs

| KPI | Definition | Formula | Target | Alert Threshold |
|-----|------------|---------|--------|-----------------|
| **Processing time (P50)** | Trigger → published artifact | `percentile(latency, 50)` | < 15 min | > 30 min |
| **Processing time (P95)** | Same | `percentile(latency, 95)` | < 2 hours | > 4 hours |
| **Extraction accuracy** | Fields matching human gold set | `correct_fields / total_fields` | ≥ 92% | < 88% |
| **Matching accuracy** | HM agrees with top-quartile match | Survey + outcome | ≥ 85% | < 80% |
| **Human review rate** | Items requiring review | `review_count / total` | ≤ 30% | > 45% |
| **Automation success rate** | Completed without human edit | `auto_publish / total` | ≥ 70% | < 60% |
| **Interview generation time** | Request → published guide | `percentile(latency, 50)` | < 5 min | > 15 min |
| **False positive rate** | Shortlisted but rejected later | `fp / shortlisted` | < 15% | > 20% |
| **False negative rate** | Rejected but HM wanted interview | `fn / rejected` | < 10% | > 15% |
| **Error rate** | Failed runs | `errors / total` | < 2% | > 5% |
| **Override rate** | Human overrides of AI | `overrides / total` | < 8% | > 12% |

---

## 3. Operational Metrics

| Metric | Purpose |
|--------|---------|
| Queue depth | Capacity planning |
| DLQ count | Reliability |
| LLM token usage | Cost management |
| OCR invocation rate | Document quality feedback to sourcing |
| Retry count | Infrastructure health |
| ATS sync failure rate | Integration health |

---

## 4. Dashboards

1. **Executive:** Time-to-fill impact, automation rate, cost per hire delta
2. **TA Operations:** Real-time queue, SLA breaches, review backlog
3. **AI Operations:** Prompt version performance, confidence distributions, error codes
4. **Compliance:** Override audit, PII access logs, bias metric trends

---

## 5. Reporting Cadence

| Report | Audience | Frequency |
|--------|----------|-----------|
| Daily ops snapshot | TA Lead | Daily |
| Weekly KPI review | HR Technology | Weekly |
| Monthly quality report | HRBP + Legal | Monthly |
| Quarterly business review | CHRO | Quarterly |

---

## 6. Baseline & Measurement Method

- **Baseline period:** 90 days pre-automation (manual process timings from ATS timestamps).
- **Gold set:** 200 annotated documents per workflow for accuracy evals.
- **HM agreement surveys:** Embedded in review console, 1-click on 20% sample.

---

*End of Metrics — {wf['id']}*
"""


def example_input_md(folder: str, wf: dict) -> str:
    return f"""# {wf['name']} — Example Input

**Workflow ID:** {wf['id']}  

---

## 1. Overview

This document provides a realistic example input for **{wf['name']}** as would be received in production.

---

## 2. Metadata Envelope

```json
{{
  "workflow_id": "{wf['id']}",
  "requisition_id": "REQ-2026-0142",
  "candidate_id": "CAN-88421",
  "workflow_run_id": "RUN-7f3a2b1c-2026-0803-001",
  "source_channel": "careers_portal",
  "priority": "standard",
  "submitted_at": "2026-08-03T09:15:00+05:30",
  "submitter_email": "talent.acquisition@contoso.com",
  "schema_version": "1.0.0",
  "source_file": {{
    "name": "Rahul_Sharma_Resume.pdf",
    "mime_type": "application/pdf",
    "size_bytes": 245760,
    "sha256": "a1b2c3d4e5f6789012345678901234567890abcdef1234567890abcdef123456"
  }}
}}
```

---

## 3. Primary Content

See `samples/` for full files. Excerpt for **{wf['short']}**:

---

### Job Description Example (WF-01)

```
Software Engineer II — Cloud Platform
Department: Engineering | Location: Bangalore (Hybrid)
Experience: 3–5 years

We are seeking a Software Engineer II to build scalable microservices on Azure...

Required Skills: Python, REST APIs, Docker, Kubernetes, SQL
Preferred: Azure, CI/CD, system design
Education: BS Computer Science or equivalent
```

---

### Resume Excerpt (WF-02/03)

```
RAHUL SHARMA
Email: rahul.sharma@email.com | Phone: +91-98765-43210 | Bangalore, India

SUMMARY
Software engineer with 4.5 years building backend services and cloud-native applications.

EXPERIENCE
Senior Software Engineer | Infosys | Jan 2022 – Present
- Designed REST microservices in Python serving 2M daily requests
- Deployed services on Azure AKS with Docker/Kubernetes
- Reduced API latency 35% via query optimization and caching

Software Engineer | TCS | Jun 2020 – Dec 2021
- Built CI/CD pipelines using Azure DevOps
- Developed SQL reporting modules for enterprise clients

EDUCATION
B.Tech Computer Science, VTU, 2020 — CGPA 8.4/10

SKILLS
Python, Java, REST, Docker, Kubernetes, Azure, PostgreSQL, Git
```

---

## 4. Input Validation Checklist

- [x] Valid file format
- [x] Required metadata present
- [x] Requisition exists and is OPEN
- [x] Consent flag true
- [x] File size within limits

---

*End of Example Input — {wf['id']}*
"""


def example_output_md(folder: str, wf: dict) -> str:
    return f"""# {wf['name']} — Example Output

**Workflow ID:** {wf['id']}  

---

## 1. Successful Output Example

```json
{{
  "status": "SUCCESS",
  "workflow_id": "{wf['id']}",
  "prompt_version": "1.0.0",
  "model_id": "gpt-4o-2026-05-01",
  "requisition_id": "REQ-2026-0142",
  "candidate_id": "CAN-88421",
  "processed_at": "2026-08-03T09:18:42Z",
  "confidence_aggregate": 0.93,
  "human_review_required": false,
  "results": {{
    "summary": "Strong match for Software Engineer II — 4.5 years Python/Azure experience aligns with core requirements.",
    "key_findings": [
      "Python and Kubernetes experience confirmed in current role",
      "Azure AKS deployment experience matches preferred qualifications",
      "Education requirement met — B.Tech CS 2020"
    ]
  }},
  "evidence": [
    {{"field": "python_experience", "snippet": "Designed REST microservices in Python", "source_location": "page 1"}},
    {{"field": "kubernetes", "snippet": "Deployed services on Azure AKS with Docker/Kubernetes", "source_location": "page 1"}}
  ],
  "flags": [],
  "artifact_path": "{folder}/REQ-2026-0142_CAN-88421_output_20260803.json"
}}
```

---

## 2. Partial Output Example

```json
{{
  "status": "PARTIAL",
  "confidence_aggregate": 0.74,
  "human_review_required": true,
  "review_reasons": ["LOW_CONFIDENCE_EMPLOYMENT_DATES"],
  "flags": ["DATE_AMBIGUITY_TCS_END"]
}}
```

---

## 3. Failed Output Example

```json
{{
  "status": "FAILED",
  "confidence_aggregate": 0.0,
  "human_review_required": true,
  "review_reasons": ["ENCRYPTED_DOCUMENT"],
  "error_code": "ERR_FILE_ENCRYPTED"
}}
```

---

## 4. Downstream Consumption

| Consumer | Uses Field |
|----------|------------|
| Next workflow | Full `results` object |
| ATS | `status`, summary, artifact link |
| Review Console | `flags`, `review_reasons`, evidence |
| Analytics | `confidence_aggregate`, latency metadata |

---

*End of Example Output — {wf['id']}*
"""


def human_review_md(folder: str, wf: dict) -> str:
    return f"""# {wf['name']} — Human Review Guide

**Workflow ID:** {wf['id']}  
**Version:** 1.0.0  

---

## 1. Automation vs Human Boundaries

| Step | Automation Level | Human Role | Never Fully Automate Because |
|------|------------------|------------|------------------------------|
| File intake & virus scan | Fully automated | Monitor exceptions | — |
| Text extraction / OCR | Automated with fallback | Manual transcription | Quality for scanned docs |
| AI {wf['short']} | Automated | Review low-confidence | Accuracy & liability |
| Schema validation | Fully automated | Override with reason | — |
| Publish to artifact store | Automated post-approval | Approve/reject gate | Accountability |
| ATS status update | Automated | Verify on failure | Candidate experience |
| Candidate rejection comms | Draft automated | **HR must approve** | Legal/compliance |
| Final hire decision | AI compiles dossier | **HM + HRBP decide** | Employment law |

---

## 2. Approval Matrix

| Action | TA Specialist | Hiring Manager | HRBP | Legal |
|--------|:-------------:|:--------------:|:----:|:-----:|
| Accept AI extraction | ✓ Approve | Inform | — | — |
| Edit extracted fields | ✓ | — | — | — |
| Approve match for interview | Recommend | ✓ Approve | Inform | — |
| Approve interview guide | Recommend | ✓ Approve | — | Consult |
| Approve offer | Prepare | Recommend | ✓ Approve | Consult |
| Send rejection | Draft | — | ✓ Approve | — |

---

## 3. When Human Review Is Mandatory

1. `confidence_aggregate` < 0.85
2. Any compliance flag present
3. First 50 runs after prompt version change
4. Candidate is internal employee or referral executive
5. HM explicitly enables "always review" on requisition
6. Discrepancy between AI output and prior ATS data

---

## 4. Review SLA

| Priority | Review SLA | Escalation |
|----------|------------|------------|
| Critical role | 2 hours | TA Lead |
| Standard | 4 hours | TA Specialist backup |
| Batch/low | 24 hours | Weekly cleanup |

---

## 5. Reviewer Instructions

1. Open Review Console task linked from email/Teams notification.
2. Compare source document (left pane) with AI output (right pane).
3. Verify evidence snippets support extracted fields.
4. Edit incorrect fields inline; system re-validates on save.
5. Select **Approve**, **Reject**, or **Escalate to Compliance**.
6. Provide reason code for any material override.

---

## 6. Why Humans Must Retain Authority

- **Legal:** Employment decisions require human accountability in most jurisdictions.
- **Ethical:** AI may encode bias; humans apply context and fairness.
- **Quality:** Edge cases (career gaps, non-linear paths) need judgment.
- **Trust:** Hiring managers must own team composition decisions.
- **Compliance:** EEOC and GDPR require explainable, auditable processes.

---

*End of Human Review Guide — {wf['id']}*
"""


WORKFLOW_FILES = {
    "Workflow_Spec.md": workflow_spec,
    "Prompt.md": prompt_md,
    "Execution_Flow.md": execution_flow_md,
    "Architecture.md": architecture_md,
    "Automation.md": automation_md,
    "Error_Handling.md": error_handling_md,
    "Validation.md": validation_md,
    "Metrics.md": metrics_md,
    "Example_Input.md": example_input_md,
    "Example_Output.md": example_output_md,
    "Human_Review.md": human_review_md,
}


PROMPT_LIBRARY = [
    "Resume Parser", "JD Parser", "JD Matcher", "Candidate Summary",
    "Interview Generator", "Interview Evaluation", "Offer Letter Generator",
    "Email Generator", "Candidate Rejection", "Knowledge Base Search",
]

SOPS = [
    ("Resume_Screening_SOP", "Resume Screening"),
    ("Interview_SOP", "Interview Process"),
    ("Hiring_SOP", "Hiring and Offer"),
    ("Prompt_Governance_SOP", "Prompt Governance"),
    ("Data_Privacy_SOP", "Data Privacy"),
    ("Knowledge_Management_SOP", "Knowledge Management"),
    ("AI_Quality_Review_SOP", "AI Quality Review"),
    ("Prompt_Testing_SOP", "Prompt Testing"),
    ("Prompt_Versioning_SOP", "Prompt Versioning"),
    ("Automation_Maintenance_SOP", "Automation Maintenance"),
]


def sop_content(title: str, name: str) -> str:
    return f"""# {name} — Standard Operating Procedure

**SOP ID:** SOP-{name.upper().replace(' ', '_')[:20]}  
**Version:** 1.0.0  
**Effective Date:** {datetime.now().strftime('%Y-%m-%d')}  
**Classification:** Internal  

---

## 1. Purpose

This Standard Operating Procedure (SOP) defines the repeatable process, roles, and controls for **{name.lower()}** within the AI Recruitment Operations platform. It ensures consistent execution, regulatory compliance, and alignment with enterprise talent acquisition policies.

---

## 2. Scope

**In scope:**
- All full-time employee requisitions processed through the AI Recruitment Operations pipeline
- TA Specialists, Recruiters, Hiring Managers, HRBPs, and Interviewers using automated workflows 01–08
- Prompt library assets and automation integrations supporting {name.lower()}

**Out of scope:**
- Contractor/contingent workforce (see Contingent Workforce SOP)
- Executive search retained firms
- Pre-offer background check vendor operations (referenced only)

---

## 3. Owner

| Role | Name/Title | Contact |
|------|------------|---------|
| **Process Owner** | Director, Talent Acquisition Operations | ta-ops@contoso.com |
| **Document Owner** | AI Operations Lead | ai-ops@contoso.com |
| **Compliance Owner** | HR Compliance Manager | hr-compliance@contoso.com |

---

## 4. Frequency

| Activity | Frequency |
|----------|-----------|
| Execute process | Per requisition / per candidate event |
| SOP compliance audit | Quarterly |
| SOP document review | Semi-annual |
| Training refresh | Annual + on role change |

---

## 5. Procedure

### 5.1 Initiation
1. Verify requisition is approved in ATS with valid headcount.
2. Confirm AI workflows enabled for department (check `docs/rollout_matrix.md`).
3. Assign TA Specialist as primary operator for requisition.

### 5.2 Execution — {name}
1. Follow workflow-specific documentation in folders `01_Job_Descriptions/` through `08_Final_Decision/`.
2. Monitor Review Console for items exceeding SLA.
3. Document manual overrides with reason codes in ATS notes.
4. Escalate compliance flags within prescribed timeframes (see Section 9).

### 5.3 Quality Checks
1. Verify AI outputs against source documents before HM presentation.
2. Confirm prompt version logged in artifact metadata.
3. Complete spot-check per AI Quality Review SOP (10% sample weekly).

### 5.4 Closure
1. Update ATS stage to reflect completed process step.
2. Archive artifacts in designated SharePoint library.
3. Emit metrics event for reporting dashboard.

---

## 6. Roles and Responsibilities

| Role | Responsibilities |
|------|------------------|
| **TA Specialist** | Operate workflows, first-line review, candidate communication coordination |
| **Recruiter** | Source candidates, manage pipeline, schedule interviews |
| **Hiring Manager** | Approve shortlists, interview guides, final hire decision |
| **HRBP** | Offer approval, policy exceptions, org alignment |
| **Interviewer** | Conduct interviews, submit timely feedback |
| **AI Operations** | Maintain prompts, automation, incident response |
| **Legal/Compliance** | Review non-standard templates, adverse impact concerns |
| **Privacy Officer** | Data subject requests, retention, cross-border transfers |

---

## 7. Approvals

| Step | Approver | Evidence Required |
|------|----------|-------------------|
| Workflow output publication (standard) | TA Specialist | Validation pass log |
| Shortlist approval | Hiring Manager | Match report + summaries |
| Interview guide | Hiring Manager | Generated question set |
| Offer release | HRBP + HM | Decision dossier |
| Rejection at final stage | HRBP | Rejection template approval |
| Prompt production deploy | AI Ops Lead + HR Compliance | Eval results |

---

## 8. Risks

| Risk | Impact | Control |
|------|--------|---------|
| Biased screening | Legal, reputational | DEI audits, blind review options |
| PII mishandling | Regulatory fines | Data Privacy SOP, encryption |
| Over-reliance on AI | Poor hires | Mandatory human gates |
| Process bypass | Audit failure | ATS-enforced stage gates |
| Stale documentation | Operational errors | Semi-annual SOP review |

---

## 9. Escalation

| Level | Trigger | Escalate To | Timeframe |
|-------|---------|-------------|-----------|
| L1 | SLA breach, single error | TA Team Lead | 4 hours |
| L2 | Compliance flag, bias concern | HR Compliance | 24 hours |
| L3 | PII incident, system outage | CISO + HR Director | Immediate |
| L4 | Legal hold, litigation | Legal Counsel | Immediate |

**Escalation channel:** `#talent-tech-incidents` Teams channel + ServiceNow ticket category `HR-AI-OPS`.

---

## 10. Related Documents

- `10_SOPs/` — all companion SOPs
- `01_Job_Descriptions/` through `08_Final_Decision/` — workflow specs
- `09_Prompt_Library/` — governed prompts
- `docs/governance_framework.md`

---

## 11. Revision History

| Version | Date | Author | Changes |
|---------|------|--------|---------|
| 1.0.0 | {datetime.now().strftime('%Y-%m-%d')} | AI Operations | Initial release |

---

*End of SOP — {name}*
"""


def prompt_library_content(name: str) -> str:
    slug = name.lower().replace(" ", "_")
    return f"""# Prompt Library — {name}

**Prompt ID:** LIB-{slug.upper()}  
**Version:** 1.0.0  
**Owner:** Prompt Engineering Lead  
**Last Reviewed:** {datetime.now().strftime('%Y-%m-%d')}  

---

## Overview

Production prompt for **{name}** in the AI Recruitment Operations platform. Used across one or more workflows; see cross-references below.

---

## Why Each Section Exists

| Section | Rationale |
|---------|-----------|
| System Prompt | Immutable guardrails for compliance, format, and behavior |
| User Prompt | Variable injection point for runtime data |
| Expected Output | Contract for validators and integrations |
| JSON Schema | Machine-readable validation |
| Validation Rules | Business logic beyond schema |
| Examples | Few-shot quality anchors |
| Edge Cases | Documented test scenarios |
| Version/Owner/Notes | Governance and audit trail |

---

## System Prompt

```
You are an enterprise recruitment AI performing {name.lower()} for Contoso Ltd.

RULES:
- Base all outputs strictly on provided source documents.
- Never infer protected characteristics or use them in decisions.
- Output valid JSON unless prose template explicitly requested.
- Include evidence snippets (max 25 words) for material claims.
- Set human_review_required=true when confidence < 0.85.
- Include prompt_version and prompt_id in every response.

COMPLIANCE: GDPR/CCPA minimization. EEO-neutral language. No discriminatory questions.
```

---

## User Prompt Template

```
Execute {name} for the following inputs.

requisition_id: {{{{requisition_id}}}}
candidate_id: {{{{candidate_id}}}}
workflow_run_id: {{{{workflow_run_id}}}}

=== PRIMARY SOURCE ===
{{{{source_content}}}}

=== OPTIONAL CONTEXT ===
{{{{context}}}}

Return output conforming to {slug}_output schema.
```

---

## Expected Output

Structured JSON with: `status`, `prompt_id`, `prompt_version`, `confidence_aggregate`, `results`, `evidence[]`, `human_review_required`, `flags[]`.

---

## JSON Schema Reference

Canonical schema: `schemas/prompt_metadata.json` (envelope) + workflow-specific extensions in `schemas/`.

---

## Validation Rules

- VR-01: confidence ≥ 0.85 for auto-publish
- VR-02: evidence fuzzy-match against source ≥ 0.8
- VR-03: no protected-class language in decision fields
- VR-04: required fields per schema

---

## Examples

See `samples/` and workflow `Example_Output.md` files for full worked examples.

---

## Edge Cases

- Empty source → FAILED
- Encrypted PDF → FAILED with ERR_FILE_ENCRYPTED
- Multi-language → process with locale tag, reduce confidence if uncertain
- Conflicting sources → flag, do not merge silently

---

## Cross-References

| Workflow | Usage |
|----------|-------|
| 01 Job Descriptions | JD Parser |
| 03 Extracted Data | Resume Parser |
| 04 Match Results | JD Matcher |
| 05 Shortlisted | Candidate Summary |
| 06 Interview Questions | Interview Generator |
| 07 Interview Feedback | Interview Evaluation |
| 08 Final Decision | Offer Letter, Email, Rejection |

---

## Version History

| Version | Date | Changes |
|---------|------|---------|
| 1.0.0 | {datetime.now().strftime('%Y-%m-%d')} | Initial production release |

---

*End of Prompt Library Entry — {name}*
"""


def generate_schemas():
    schemas = {
        "resume.schema.json": {
            "$schema": "https://json-schema.org/draft/2020-12/schema",
            "$id": "https://recruitment.contoso.com/schemas/resume.json",
            "title": "Resume",
            "type": "object",
            "required": ["candidate_id", "personal_info", "experience", "education", "skills"],
            "properties": {
                "candidate_id": {"type": "string", "pattern": "^CAN-[0-9]+$"},
                "personal_info": {
                    "type": "object",
                    "required": ["full_name", "email"],
                    "properties": {
                        "full_name": {"type": "string"},
                        "email": {"type": "string", "format": "email"},
                        "phone": {"type": "string"},
                        "location": {"type": "string"},
                        "linkedin_url": {"type": "string", "format": "uri"}
                    }
                },
                "summary": {"type": "string"},
                "experience": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "required": ["company", "title", "start_date"],
                        "properties": {
                            "company": {"type": "string"},
                            "title": {"type": "string"},
                            "start_date": {"type": "string", "format": "date"},
                            "end_date": {"type": ["string", "null"], "format": "date"},
                            "is_current": {"type": "boolean"},
                            "description": {"type": "string"},
                            "skills_used": {"type": "array", "items": {"type": "string"}}
                        }
                    }
                },
                "education": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "required": ["institution", "degree"],
                        "properties": {
                            "institution": {"type": "string"},
                            "degree": {"type": "string"},
                            "field_of_study": {"type": "string"},
                            "graduation_date": {"type": "string", "format": "date"},
                            "gpa": {"type": "string"}
                        }
                    }
                },
                "skills": {"type": "array", "items": {"type": "string"}},
                "certifications": {"type": "array", "items": {"type": "string"}},
                "extraction_metadata": {
                    "type": "object",
                    "properties": {
                        "confidence_aggregate": {"type": "number", "minimum": 0, "maximum": 1},
                        "prompt_version": {"type": "string"},
                        "extracted_at": {"type": "string", "format": "date-time"}
                    }
                }
            }
        },
        "job_description.schema.json": {
            "$schema": "https://json-schema.org/draft/2020-12/schema",
            "$id": "https://recruitment.contoso.com/schemas/job_description.json",
            "title": "Job Description",
            "type": "object",
            "required": ["requisition_id", "title", "department", "required_skills", "experience_years"],
            "properties": {
                "requisition_id": {"type": "string", "pattern": "^REQ-[0-9]{4}-[0-9]+$"},
                "title": {"type": "string"},
                "level": {"type": "string", "enum": ["Entry", "Mid", "Senior", "Staff", "Principal", "Director"]},
                "department": {"type": "string"},
                "location": {"type": "string"},
                "employment_type": {"type": "string", "enum": ["Full-time", "Part-time", "Contract"]},
                "experience_years": {"type": "object", "properties": {"min": {"type": "integer"}, "max": {"type": "integer"}}},
                "required_skills": {"type": "array", "items": {"type": "string"}},
                "preferred_skills": {"type": "array", "items": {"type": "string"}},
                "education_requirements": {"type": "string"},
                "responsibilities": {"type": "array", "items": {"type": "string"}},
                "compensation_band": {"type": "object", "properties": {"min": {"type": "number"}, "max": {"type": "number"}, "currency": {"type": "string"}}},
                "eeo_statement": {"type": "string"},
                "status": {"type": "string", "enum": ["DRAFT", "OPEN", "ON_HOLD", "CLOSED", "FILLED"]}
            }
        },
        "candidate_profile.schema.json": {
            "$schema": "https://json-schema.org/draft/2020-12/schema",
            "$id": "https://recruitment.contoso.com/schemas/candidate_profile.json",
            "title": "Candidate Profile",
            "allOf": [{"$ref": "resume.schema.json"}],
            "properties": {
                "applications": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "requisition_id": {"type": "string"},
                            "applied_at": {"type": "string", "format": "date-time"},
                            "source": {"type": "string"},
                            "stage": {"type": "string"}
                        }
                    }
                },
                "match_scores": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "requisition_id": {"type": "string"},
                            "overall_score": {"type": "number", "minimum": 0, "maximum": 1},
                            "skill_score": {"type": "number"},
                            "experience_score": {"type": "number"},
                            "education_score": {"type": "number"}
                        }
                    }
                }
            }
        },
        "interview_feedback.schema.json": {
            "$schema": "https://json-schema.org/draft/2020-12/schema",
            "$id": "https://recruitment.contoso.com/schemas/interview_feedback.json",
            "title": "Interview Feedback",
            "type": "object",
            "required": ["feedback_id", "candidate_id", "requisition_id", "interviewer", "stage", "recommendation"],
            "properties": {
                "feedback_id": {"type": "string"},
                "candidate_id": {"type": "string"},
                "requisition_id": {"type": "string"},
                "interviewer": {"type": "object", "properties": {"name": {"type": "string"}, "email": {"type": "string", "format": "email"}, "role": {"type": "string"}}},
                "stage": {"type": "string", "enum": ["Phone Screen", "Technical", "Behavioral", "Panel", "Final"]},
                "interview_date": {"type": "string", "format": "date-time"},
                "competency_ratings": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "required": ["competency", "rating"],
                        "properties": {
                            "competency": {"type": "string"},
                            "rating": {"type": "integer", "minimum": 1, "maximum": 5},
                            "notes": {"type": "string"}
                        }
                    }
                },
                "strengths": {"type": "array", "items": {"type": "string"}},
                "concerns": {"type": "array", "items": {"type": "string"}},
                "recommendation": {"type": "string", "enum": ["Strong Hire", "Hire", "No Hire", "Strong No Hire"]},
                "summary": {"type": "string"},
                "confidence_aggregate": {"type": "number", "minimum": 0, "maximum": 1}
            }
        },
        "hiring_recommendation.schema.json": {
            "$schema": "https://json-schema.org/draft/2020-12/schema",
            "$id": "https://recruitment.contoso.com/schemas/hiring_recommendation.json",
            "title": "Hiring Recommendation",
            "type": "object",
            "required": ["recommendation_id", "candidate_id", "requisition_id", "decision", "approvers"],
            "properties": {
                "recommendation_id": {"type": "string"},
                "candidate_id": {"type": "string"},
                "requisition_id": {"type": "string"},
                "decision": {"type": "string", "enum": ["OFFER", "REJECT", "HOLD", "ADDITIONAL_INTERVIEW"]},
                "overall_match_score": {"type": "number"},
                "interview_consensus": {"type": "string"},
                "compensation_proposal": {
                    "type": "object",
                    "properties": {
                        "base_salary": {"type": "number"},
                        "currency": {"type": "string"},
                        "equity": {"type": "string"},
                        "signing_bonus": {"type": "number"},
                        "within_band": {"type": "boolean"}
                    }
                },
                "approvers": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "role": {"type": "string"},
                            "name": {"type": "string"},
                            "approved_at": {"type": "string", "format": "date-time"},
                            "status": {"type": "string", "enum": ["PENDING", "APPROVED", "REJECTED"]}
                        }
                    }
                },
                "audit_trail": {"type": "array", "items": {"type": "object"}},
                "generated_at": {"type": "string", "format": "date-time"}
            }
        },
        "prompt_metadata.schema.json": {
            "$schema": "https://json-schema.org/draft/2020-12/schema",
            "$id": "https://recruitment.contoso.com/schemas/prompt_metadata.json",
            "title": "Prompt Metadata",
            "type": "object",
            "required": ["prompt_id", "prompt_version", "owner", "model_compatibility"],
            "properties": {
                "prompt_id": {"type": "string"},
                "prompt_version": {"type": "string", "pattern": "^[0-9]+\\.[0-9]+\\.[0-9]+$"},
                "name": {"type": "string"},
                "description": {"type": "string"},
                "owner": {"type": "string"},
                "workflow_ids": {"type": "array", "items": {"type": "string"}},
                "model_compatibility": {"type": "array", "items": {"type": "string"}},
                "temperature": {"type": "number", "minimum": 0, "maximum": 2},
                "max_tokens": {"type": "integer"},
                "created_at": {"type": "string", "format": "date-time"},
                "approved_by": {"type": "string"},
                "eval_score": {"type": "number"},
                "status": {"type": "string", "enum": ["DRAFT", "TESTING", "APPROVED", "DEPRECATED"]}
            }
        },
    }
    for fname, schema in schemas.items():
        write(ROOT / "schemas" / fname, json.dumps(schema, indent=2))


def generate_diagrams():
    diagrams = {
        "01_overall_architecture.mmd": """flowchart TB
    subgraph External
        C[Candidates]
        HM[Hiring Managers]
        INT[Interviewers]
    end
    subgraph Intake
        PORTAL[Careers Portal]
        EMAIL[Email Intake]
        SP[SharePoint]
    end
    subgraph AI_Recruitment_Platform
        ORCH[Workflow Orchestrator]
        LLM[LLM Gateway]
        VAL[Validation Engine]
        RC[Review Console]
        PL[Prompt Library]
    end
    subgraph Storage
        ART[Artifact Store 01-08]
        SCH[Schema Registry]
        DLQ[Dead Letter Queue]
    end
    subgraph Enterprise
        ATS[Applicant Tracking System]
        BI[Power BI Analytics]
        SIEM[SIEM Audit]
    end
    C --> PORTAL
    C --> EMAIL
    HM --> SP
    PORTAL --> ORCH
    EMAIL --> ORCH
    SP --> ORCH
    PL --> LLM
    ORCH --> LLM
    ORCH --> VAL
    LLM --> VAL
    VAL --> RC
    HM --> RC
    INT --> RC
    RC --> ART
    VAL --> ART
    ORCH --> ATS
    ART --> BI
    ORCH --> SIEM
    VAL --> DLQ
""",
        "02_folder_structure.mmd": """flowchart TD
    ROOT[AI-Recruitment-System] --> W1[01_Job_Descriptions]
    ROOT --> W2[02_Incoming_Resumes]
    ROOT --> W3[03_Extracted_Data]
    ROOT --> W4[04_Match_Results]
    ROOT --> W5[05_Shortlisted]
    ROOT --> W6[06_Interview_Questions]
    ROOT --> W7[07_Interview_Feedback]
    ROOT --> W8[08_Final_Decision]
    ROOT --> PL[09_Prompt_Library]
    ROOT --> SOP[10_SOPs]
    ROOT --> DOCS[docs]
    ROOT --> DIAG[diagrams]
    ROOT --> SCH[schemas]
    ROOT --> SAMP[samples]
    W1 --> F1[11 workflow files each]
    W2 --> F1
    W3 --> F1
    W4 --> F1
    W5 --> F1
    W6 --> F1
    W7 --> F1
    W8 --> F1
""",
        "03_recruitment_workflow.mmd": """flowchart LR
    JD[01 JD Parsing] --> INTAKE[02 Resume Intake]
    INTAKE --> EXTRACT[03 Data Extraction]
    EXTRACT --> MATCH[04 JD Matching]
    MATCH --> SHORT[05 Shortlisting]
    SHORT --> IQ[06 Interview Questions]
    IQ --> INTERVIEW[Human Interview]
    INTERVIEW --> FB[07 Feedback Processing]
    FB --> DECIDE[08 Final Decision]
    DECIDE --> OFFER[Offer / Reject]
""",
        "04_prompt_flow.mmd": """flowchart TD
    REG[Prompt Registry] --> SYS[System Prompt]
    REG --> USER[User Prompt Template]
    INPUT[Runtime Input] --> USER
    SYS --> LLM[LLM API]
    USER --> LLM
    LLM --> RAW[Raw Response]
    RAW --> PARSE[JSON Parse]
    PARSE --> SCHEMA[Schema Validation]
    SCHEMA --> BIZ[Business Rules]
    BIZ --> CONF{Confidence OK?}
    CONF -->|Yes| OUT[Published Output]
    CONF -->|No| REV[Human Review]
    REV --> OUT
""",
        "05_automation_flow.mmd": """flowchart LR
    TRIG[Webhook/Email/Schedule] --> N8N[n8n Orchestrator]
    N8N --> PY[Python Workers]
    PY --> OCR[OCR Service]
    PY --> AI[Azure OpenAI]
    N8N --> SP[SharePoint]
    N8N --> ATS[ATS API]
    N8N --> SLACK[Teams Notify]
    PY --> DLQ[Dead Letter Queue]
""",
        "06_human_approval_flow.mmd": """flowchart TD
    AI[AI Output] --> AUTO{Auto-publish eligible?}
    AUTO -->|Yes| PUB[Publish Artifact]
    AUTO -->|No| TA[TA Specialist Review]
    TA --> HM{HM Approval needed?}
    HM -->|Yes| HMA[Hiring Manager Approval]
    HM -->|No| PUB
    HMA --> HRBP{Offer/Reject?}
    HRBP -->|Offer| OFF[HRBP Offer Approval]
    HRBP -->|Reject| REJ[HRBP Rejection Approval]
    OFF --> PUB
    REJ --> PUB
""",
        "07_data_flow.mmd": """flowchart LR
    PDF[PDF/DOCX] --> OCR[Text Extraction]
    OCR --> JSON1[Structured JSON]
    JD[JD JSON] --> MATCH[Matcher]
    JSON1 --> MATCH
    MATCH --> REPORT[Match Report]
    REPORT --> SHORT[Shortlist JSON]
    SHORT --> IQ[Question Doc]
    IQ --> FB[Feedback JSON]
    FB --> DEC[Decision Package]
    DEC --> ATS[(ATS)]
""",
    }
    for fname, content in diagrams.items():
        write(ROOT / "diagrams" / fname, content)


def generate_readme():
    content = f"""# AI Recruitment Operations Platform

**Version:** 1.0.0  
**Last Updated:** {datetime.now().strftime('%Y-%m-%d')}  
**Maintained By:** AI Operations — Talent Technology  
**Classification:** Internal — Enterprise Documentation  

---

## Project Overview

The **AI Recruitment Operations Platform** is a comprehensive documentation and automation blueprint for enterprise talent acquisition. It defines how an AI Operations Specialist designs, documents, and automates the end-to-end hiring workflow—from job description parsing through final hiring decisions.

This repository is **documentation-first**: it does not ship application code, but provides production-quality specifications sufficient for an engineering team to build, deploy, and operate the system.

### Business Value

| Capability | Impact |
|------------|--------|
| Standardized JD parsing | Consistent job postings, better applicant fit |
| Automated resume extraction | 60%+ reduction in manual ATS data entry |
| AI-powered matching | Faster shortlisting with explainable scores |
| Interview question generation | Calibrated, compliant interview guides |
| Feedback synthesis | Reduced decision latency post-interview |
| Governed prompts & SOPs | Audit-ready, compliant AI operations |

---

## Architecture

The platform follows an **event-driven, schema-first, human-in-the-loop** architecture.

```mermaid
flowchart TB
    subgraph Intake
        A[Careers Portal / Email / SharePoint]
    end
    subgraph Platform
        B[Orchestrator n8n + Python]
        C[LLM Gateway + Prompt Library]
        D[Validation Engine]
        E[Review Console]
    end
    subgraph Output
        F[Artifact Store 01-08]
        G[ATS Integration]
        H[Analytics]
    end
    A --> B --> C --> D
    D --> E
    D --> F
    E --> F
    F --> G
    F --> H
```

See `diagrams/01_overall_architecture.mmd` for the complete diagram.

---

## Folder Structure

```
AI-Recruitment-System/
├── 01_Job_Descriptions/     # WF-01: JD parsing and management
├── 02_Incoming_Resumes/     # WF-02: Resume intake and routing
├── 03_Extracted_Data/       # WF-03: Resume data extraction
├── 04_Match_Results/        # WF-04: JD-candidate matching
├── 05_Shortlisted/          # WF-05: Shortlist approval
├── 06_Interview_Questions/  # WF-06: Interview guide generation
├── 07_Interview_Feedback/   # WF-07: Feedback processing
├── 08_Final_Decision/       # WF-08: Hire/reject decision pack
├── 09_Prompt_Library/       # Reusable governed prompts
├── 10_SOPs/                 # Standard operating procedures
├── docs/                    # Governance and reference docs
├── diagrams/                # Mermaid architecture diagrams
├── schemas/                 # JSON Schema definitions
├── samples/                 # Realistic example artifacts
└── README.md
```

Each workflow folder (01–08) contains **11 documents**:
`Workflow_Spec.md`, `Prompt.md`, `Execution_Flow.md`, `Architecture.md`, `Automation.md`, `Error_Handling.md`, `Validation.md`, `Metrics.md`, `Example_Input.md`, `Example_Output.md`, `Human_Review.md`.

---

## Workflow Summary

| # | Workflow | Trigger | Primary Output | Human Gate |
|---|----------|---------|----------------|------------|
| 01 | Job Description Parsing | HM submits JD | Structured JD JSON | TA review |
| 02 | Resume Intake | Candidate applies | Intake receipt + routing | Exception only |
| 03 | Data Extraction | Resume queued | Candidate profile JSON | Low confidence |
| 04 | JD Matching | Profile + JD ready | Match score report | Optional |
| 05 | Shortlisting | Match published | Approved shortlist | **HM approval** |
| 06 | Interview Questions | Interview scheduled | Question guide | **HM approval** |
| 07 | Feedback Processing | Interviewer submits | Feedback synthesis | Conflict escalation |
| 08 | Final Decision | All stages complete | Offer/reject package | **HM + HRBP** |

---

## Technology Stack

| Layer | Technology |
|-------|------------|
| LLM | Azure OpenAI (GPT-4o), Claude 3.5 (eval) |
| Orchestration | n8n (self-hosted), Python 3.11+ workers |
| Validation | JSON Schema, custom business rules |
| Storage | SharePoint Online, Azure Blob (archive) |
| ATS | Workday Recruiting / Greenhouse (API) |
| Identity | Azure AD SSO |
| Monitoring | Datadog, Power BI, SIEM |
| OCR | Azure Document Intelligence |

---

## Automation Stack

| Tool | Role |
|------|------|
| **Cursor** | Prompt development, eval scripts, documentation |
| **Python** | OCR, validation, ATS SDK, batch processing |
| **n8n** | Primary workflow orchestration, retries, webhooks |
| **Zapier** | Lightweight email → storage connectors |
| **Make.com** | Microsoft 365 + OpenAI integration scenarios |
| **REST APIs** | LLM gateway, ATS, schema registry |
| **Webhooks** | ATS events, upstream workflow handoffs |
| **Google Drive** | Regional team intake (optional) |
| **SharePoint** | Primary document library for artifacts |
| **Email** | Application parsing, notifications |

---

## Prompt Library

Located in `09_Prompt_Library/`:

1. Resume Parser
2. JD Parser
3. JD Matcher
4. Candidate Summary
5. Interview Generator
6. Interview Evaluation
7. Offer Letter Generator
8. Email Generator
9. Candidate Rejection
10. Knowledge Base Search

All prompts are versioned, owned, and governed per `10_SOPs/Prompt_Governance_SOP.md`.

---

## Documentation Index

| Category | Location | Count |
|----------|----------|-------|
| Workflow specs | `01_` – `08_` | 88 files |
| Prompt library | `09_Prompt_Library/` | 10 files |
| SOPs | `10_SOPs/` | 10 files |
| JSON schemas | `schemas/` | 6 files |
| Diagrams | `diagrams/` | 7 files |
| Sample artifacts | `samples/` | 9 files |
| Governance docs | `docs/` | 3 files |

---

## Future Roadmap

### Q3 2026
- [ ] Workday Recruiting bi-directional API integration spec
- [ ] Multimodal portfolio/GitHub enrichment with consent flow
- [ ] Hindi and Spanish resume parsing locale packs

### Q4 2026
- [ ] Active learning pipeline from reviewer corrections
- [ ] Predictive time-to-fill analytics model
- [ ] Candidate self-scheduling integration

### 2027
- [ ] Internal mobility workflow extension
- [ ] DEI adverse impact monitoring dashboard
- [ ] Full SOC 2 Type II audit package for AI hiring ops

---

## Getting Started

1. Read this README and `docs/governance_framework.md`.
2. Review workflow `01_Job_Descriptions/Workflow_Spec.md` as the entry point.
3. Examine `samples/` for realistic input/output artifacts.
4. Follow `10_SOPs/` for operational procedures.
5. Use `diagrams/` and `schemas/` for implementation planning.

---

## Contact

**AI Operations Lead:** ai-ops@contoso.com  
**Talent Acquisition Operations:** ta-ops@contoso.com  
**HR Compliance:** hr-compliance@contoso.com  

---

*© 2026 Contoso Ltd. Internal Use Only.*
"""
    write(ROOT / "README.md", content)


def generate_docs():
    write(ROOT / "docs" / "governance_framework.md", """# AI Recruitment Operations — Governance Framework

## 1. Purpose
Defines policies for prompt lifecycle, data handling, human oversight, and audit requirements.

## 2. Principles
- Human accountability for all employment decisions
- Schema-first AI outputs with validation gates
- Versioned prompts with eval-before-deploy
- PII minimization and 7-year retention
- Explainability via evidence citations

## 3. RACI Summary
See individual workflow Human_Review.md files and 10_SOPs/.

## 4. Change Management
All prompt changes require: eval run, HR Compliance sign-off, staged rollout (dev→staging→prod).

## 5. Audit Requirements
Every artifact logs: prompt_version, model_id, workflow_run_id, reviewer_id (if applicable).
""")
    write(ROOT / "docs" / "rollout_matrix.md", """# Department Rollout Matrix

| Department | WF-01 | WF-02 | WF-03 | WF-04 | WF-05 | WF-06 | WF-07 | WF-08 |
|------------|-------|-------|-------|-------|-------|-------|-------|-------|
| Engineering | ✓ | ✓ | ✓ | ✓ | ✓ | ✓ | ✓ | ✓ |
| Product | ✓ | ✓ | ✓ | ✓ | ✓ | ✓ | ✓ | ✓ |
| Sales | ✓ | ✓ | ✓ | ✓ | Partial | Partial | ✓ | ✓ |
| Operations | ✓ | ✓ | ✓ | ✓ | ✓ | ✓ | ✓ | ✓ |

*Partial = interview question AI assist only; HM writes final questions.*
""")
    write(ROOT / "docs" / "implementation_guide.md", """# Implementation Guide for Engineering Teams

## Phase 1 — Foundation (Weeks 1–4)
1. Deploy schema registry and artifact storage (SharePoint + blob).
2. Implement Python validation worker with jsonschema.
3. Configure n8n webhook intake for WF-02.

## Phase 2 — Core AI Workflows (Weeks 5–10)
1. Deploy prompts from 09_Prompt_Library/ via LLM gateway.
2. Build Review Console MVP (accept/reject/edit).
3. Integrate ATS stage updates.

## Phase 3 — Full Pipeline (Weeks 11–16)
1. Connect workflows 01–08 via event bus.
2. Implement DLQ, retry, monitoring dashboards.
3. Run parallel with manual process for 30-day validation.

## Phase 4 — Production (Week 17+)
1. Cutover per rollout_matrix.md by department.
2. Enable KPI dashboards and compliance reporting.
""")


def generate_sample_json():
    jd = {
        "requisition_id": "REQ-2026-0142",
        "title": "Software Engineer II — Cloud Platform",
        "level": "Mid",
        "department": "Engineering",
        "location": "Bangalore, India (Hybrid)",
        "employment_type": "Full-time",
        "experience_years": {"min": 3, "max": 5},
        "required_skills": ["Python", "REST APIs", "Docker", "Kubernetes", "SQL"],
        "preferred_skills": ["Azure", "CI/CD", "System Design", "Microservices"],
        "education_requirements": "Bachelor's degree in Computer Science or equivalent practical experience",
        "responsibilities": [
            "Design and implement scalable REST microservices",
            "Deploy and operate containerized workloads on Azure AKS",
            "Collaborate with product and platform teams on architecture decisions",
            "Participate in code reviews and mentor junior engineers",
            "Contribute to CI/CD pipeline improvements"
        ],
        "compensation_band": {"min": 1800000, "max": 2800000, "currency": "INR"},
        "eeo_statement": "Contoso is an equal opportunity employer. All qualified applicants will receive consideration without regard to race, color, religion, sex, national origin, disability, or veteran status.",
        "status": "OPEN"
    }
    write(ROOT / "samples" / "Software_Engineer.json", json.dumps(jd, indent=2))

    resume = {
        "candidate_id": "CAN-88421",
        "personal_info": {
            "full_name": "Rahul Sharma",
            "email": "rahul.sharma@email.com",
            "phone": "+91-98765-43210",
            "location": "Bangalore, India",
            "linkedin_url": "https://linkedin.com/in/rahulsharma-dev"
        },
        "summary": "Software engineer with 4.5 years of experience building backend services and cloud-native applications on Azure. Strong Python and Kubernetes skills with proven track record of performance optimization.",
        "experience": [
            {
                "company": "Infosys Limited",
                "title": "Senior Software Engineer",
                "start_date": "2022-01-15",
                "end_date": None,
                "is_current": True,
                "description": "Designed REST microservices in Python serving 2M daily requests. Deployed services on Azure AKS with Docker/Kubernetes. Reduced API latency 35% via query optimization and Redis caching.",
                "skills_used": ["Python", "REST", "Docker", "Kubernetes", "Azure", "PostgreSQL", "Redis"]
            },
            {
                "company": "Tata Consultancy Services",
                "title": "Software Engineer",
                "start_date": "2020-06-01",
                "end_date": "2021-12-31",
                "is_current": False,
                "description": "Built CI/CD pipelines using Azure DevOps. Developed SQL reporting modules for enterprise banking clients.",
                "skills_used": ["Python", "Azure DevOps", "SQL", "Java"]
            }
        ],
        "education": [
            {
                "institution": "Visvesvaraya Technological University",
                "degree": "B.Tech",
                "field_of_study": "Computer Science",
                "graduation_date": "2020-05-15",
                "gpa": "8.4/10"
            }
        ],
        "skills": ["Python", "Java", "REST APIs", "Docker", "Kubernetes", "Azure", "PostgreSQL", "Git", "CI/CD", "Microservices"],
        "certifications": ["AZ-900: Azure Fundamentals", "CKA: Certified Kubernetes Administrator"],
        "extraction_metadata": {
            "confidence_aggregate": 0.94,
            "prompt_version": "1.0.0",
            "extracted_at": "2026-08-03T09:18:42Z"
        }
    }
    write(ROOT / "samples" / "Rahul_Sharma.json", json.dumps(resume, indent=2))


def generate_sample_pdfs_and_docx():
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    except ImportError:
        print("reportlab not installed - skipping PDF generation")
        return

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=16, spaceAfter=12)
    heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], fontSize=12, spaceAfter=6)
    body_style = styles['Normal']

    def make_pdf(path, title, sections):
        doc = SimpleDocTemplate(str(path), pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
        story = [Paragraph(title, title_style), Spacer(1, 12)]
        for heading, paragraphs in sections:
            story.append(Paragraph(heading, heading_style))
            for p in paragraphs:
                story.append(Paragraph(p, body_style))
                story.append(Spacer(1, 6))
            story.append(Spacer(1, 12))
        doc.build(story)

    jd_sections = [
        ("Position Overview", [
            "Contoso Ltd is seeking a Software Engineer II to join our Cloud Platform team in Bangalore.",
            "This hybrid role requires 3 days per week in office.",
        ]),
        ("Required Qualifications", [
            "3-5 years software engineering experience",
            "Strong Python and REST API development",
            "Experience with Docker and Kubernetes",
            "SQL database proficiency",
            "Bachelor's in CS or equivalent",
        ]),
        ("Preferred Qualifications", [
            "Azure cloud platform experience",
            "CI/CD pipeline design",
            "System design for distributed systems",
        ]),
        ("Key Responsibilities", [
            "Design and implement scalable REST microservices",
            "Deploy containerized workloads on Azure AKS",
            "Participate in code reviews and mentoring",
            "Improve CI/CD pipelines and observability",
        ]),
        ("Compensation", ["INR 18,00,000 - 28,00,000 per annum based on experience"]),
        ("EEO Statement", ["Contoso is an equal opportunity employer."]),
    ]
    make_pdf(ROOT / "samples" / "Software_Engineer.pdf", "Software Engineer II — Cloud Platform", jd_sections)

    resume_sections = [
        ("RAHUL SHARMA", [
            "Email: rahul.sharma@email.com | Phone: +91-98765-43210",
            "Location: Bangalore, India | LinkedIn: linkedin.com/in/rahulsharma-dev",
        ]),
        ("PROFESSIONAL SUMMARY", [
            "Software engineer with 4.5 years building backend services and cloud-native applications. Expertise in Python, Kubernetes, and Azure.",
        ]),
        ("EXPERIENCE", [
            "<b>Senior Software Engineer | Infosys Limited | Jan 2022 – Present</b>",
            "• Designed REST microservices in Python serving 2M daily requests",
            "• Deployed services on Azure AKS with Docker/Kubernetes",
            "• Reduced API latency 35% via query optimization and Redis caching",
            "<b>Software Engineer | TCS | Jun 2020 – Dec 2021</b>",
            "• Built CI/CD pipelines using Azure DevOps",
            "• Developed SQL reporting modules for banking clients",
        ]),
        ("EDUCATION", [
            "B.Tech Computer Science, VTU, 2020 — CGPA 8.4/10",
        ]),
        ("SKILLS", [
            "Python, Java, REST APIs, Docker, Kubernetes, Azure, PostgreSQL, Git, CI/CD",
        ]),
        ("CERTIFICATIONS", [
            "AZ-900 Azure Fundamentals | CKA Certified Kubernetes Administrator",
        ]),
    ]
    make_pdf(ROOT / "samples" / "Rahul_Sharma_Resume.pdf", "Rahul Sharma — Resume", resume_sections)

    match_sections = [
        ("Match Report", [
            "Requisition: REQ-2026-0142 — Software Engineer II",
            "Candidate: CAN-88421 — Rahul Sharma",
            "Overall Match Score: 87/100",
        ]),
        ("Skills Alignment", [
            "Python: STRONG MATCH — 4.5 years production experience",
            "Kubernetes: STRONG MATCH — Azure AKS deployment experience",
            "REST APIs: STRONG MATCH — Microservices architecture",
            "Azure: MATCH — Current role uses Azure platform",
            "System Design: PARTIAL — Limited explicit evidence",
        ]),
        ("Experience Alignment", [
            "Required: 3-5 years | Candidate: 4.5 years — MATCH",
            "Domain: Cloud platform — STRONG ALIGNMENT",
        ]),
        ("Recommendation", [
            "RECOMMEND SHORTLIST for technical interview.",
            "Gap: System design depth — probe in technical round.",
        ]),
    ]
    make_pdf(ROOT / "samples" / "Match_Report.pdf", "Candidate Match Report", match_sections)

    hire_sections = [
        ("Final Hiring Report", [
            "Requisition: REQ-2026-0142 | Candidate: Rahul Sharma (CAN-88421)",
            "Decision: OFFER RECOMMENDED",
            "Date: August 3, 2026",
        ]),
        ("Summary", [
            "Rahul Sharma demonstrated strong technical competency across all interview stages.",
            "Unanimous Hire recommendation from technical panel. Behavioral interview: Strong Hire.",
        ]),
        ("Interview Scores", [
            "Phone Screen: Hire (4.2/5 avg)",
            "Technical Round: Strong Hire (4.5/5 avg)",
            "Behavioral Round: Strong Hire (4.3/5 avg)",
            "Hiring Manager Round: Hire (4.4/5 avg)",
        ]),
        ("Compensation Proposal", [
            "Base Salary: INR 24,00,000 | Within approved band",
            "Sign-on Bonus: INR 2,00,000 | Standard new hire",
            "Start Date: September 1, 2026",
        ]),
        ("Approvals", [
            "Hiring Manager: Approved — Priya Mehta — 2026-08-02",
            "HRBP: Approved — Ankit Verma — 2026-08-03",
        ]),
    ]
    make_pdf(ROOT / "samples" / "Final_Hiring_Report.pdf", "Final Hiring Decision Report", hire_sections)

    summary_sections = [
        ("Candidate Summary — Rahul Sharma", [
            "For: Software Engineer II — Cloud Platform (REQ-2026-0142)",
        ]),
        ("Overview", [
            "4.5-year software engineer with strong Python and Azure Kubernetes experience.",
            "Currently Senior Software Engineer at Infosys building high-throughput microservices.",
        ]),
        ("Key Strengths", [
            "Production Python microservices at scale (2M daily requests)",
            "Hands-on Azure AKS and Docker/Kubernetes operations",
            "Performance optimization track record (35% latency reduction)",
            "CKA certification demonstrates platform depth",
        ]),
        ("Areas to Probe", [
            "System design for distributed systems at Contoso scale",
            "Cross-team collaboration examples",
        ]),
        ("Match Score: 87/100 — Recommended for Interview", []),
    ]
    make_pdf(ROOT / "samples" / "Candidate_Summary.pdf", "Candidate Summary", summary_sections)

    try:
        from docx import Document
        from docx.shared import Pt, Inches

        iq = Document()
        iq.add_heading("Interview Questions — Rahul Sharma", 0)
        iq.add_paragraph("Requisition: REQ-2026-0142 | Stage: Technical Round | Duration: 60 min")
        iq.add_heading("Competency: Python & Backend Development", level=1)
        for q in [
            "Describe a REST microservice you designed at Infosys. What were the key API design decisions?",
            "How did you achieve 35% latency reduction? Walk through your optimization process.",
            "Follow-up: What trade-offs did you consider between caching and data consistency?",
        ]:
            iq.add_paragraph(q, style='List Number')
        iq.add_heading("Competency: Cloud & DevOps", level=1)
        for q in [
            "Explain your Azure AKS deployment workflow from code commit to production.",
            "Describe a production incident with Kubernetes and how you resolved it.",
            "Follow-up: How do you handle secrets management in AKS?",
        ]:
            iq.add_paragraph(q, style='List Number')
        iq.add_heading("Competency: System Design", level=1)
        for q in [
            "Design a URL shortening service handling 10K requests/second.",
            "Follow-up: How would you migrate this from single-region to multi-region?",
        ]:
            iq.add_paragraph(q, style='List Number')
        iq.add_heading("Scoring Rubric", level=1)
        iq.add_paragraph("1 = Does not meet | 3 = Meets | 5 = Exceeds expectations")
        iq.save(str(ROOT / "samples" / "Interview_Questions.docx"))

        fb = Document()
        fb.add_heading("Interview Feedback — Rahul Sharma", 0)
        fb.add_paragraph("Interviewer: Vikram Singh, Staff Engineer | Date: 2026-07-28 | Stage: Technical")
        fb.add_heading("Competency Ratings", level=1)
        ratings = [
            ("Python & Backend", 5, "Excellent microservices explanation with concrete metrics"),
            ("Cloud & DevOps", 4, "Strong AKS experience; minor gap in multi-region DR"),
            ("System Design", 4, "Good URL shortener design; needed hint on sharding"),
            ("Communication", 5, "Clear, structured responses with appropriate depth"),
        ]
        for comp, rating, notes in ratings:
            fb.add_paragraph(f"{comp}: {rating}/5 — {notes}")
        fb.add_heading("Strengths", level=1)
        for s in ["Deep Python production experience", "Strong Kubernetes operational knowledge", "Quantified impact in previous roles"]:
            fb.add_paragraph(s, style='List Bullet')
        fb.add_heading("Concerns", level=1)
        fb.add_paragraph("Limited exposure to multi-region architecture — addressable with onboarding", style='List Bullet')
        fb.add_heading("Recommendation", level=1)
        fb.add_paragraph("STRONG HIRE — Recommend proceeding to behavioral round.")
        fb.save(str(ROOT / "samples" / "Interview_Feedback.docx"))
    except ImportError:
        print("python-docx not installed - skipping DOCX generation")


def main():
    ROOT.mkdir(parents=True, exist_ok=True)
    count = 0

    for folder, wf in WORKFLOWS.items():
        for fname, gen_fn in WORKFLOW_FILES.items():
            write(ROOT / folder / fname, gen_fn(folder, wf))
            count += 1

    for name in PROMPT_LIBRARY:
        slug = name.lower().replace(" ", "_")
        write(ROOT / "09_Prompt_Library" / f"{slug}.md", prompt_library_content(name))
        count += 1

    for fname, title in SOPS:
        write(ROOT / "10_SOPs" / f"{fname}.md", sop_content(title, title))
        count += 1

    generate_schemas()
    count += 6

    generate_diagrams()
    count += 7

    generate_readme()
    count += 1

    generate_docs()
    count += 3

    generate_sample_json()
    count += 2

    generate_sample_pdfs_and_docx()
    count += 7  # approximate for samples

    total = sum(1 for _ in ROOT.rglob("*") if _.is_file() and _.name != "_generate_project.py")
    print(f"Generated project at {ROOT}")
    print(f"Total files (excluding generator): {total}")


if __name__ == "__main__":
    main()
