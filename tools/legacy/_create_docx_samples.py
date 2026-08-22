# =============================================================================
# SUPERSEDED - DO NOT RUN
#
# This script overwrites the ENTIRE project tree from a hardcoded path
# (C:\AI-Recruitment-System). Running it destroys every hand-edited
# document, with no recovery path outside git.
#
# As of the initial commit the documentation is hand-maintained. This file is
# kept only as a record of how the original documentation was produced.
#
# See CLAUDE.md -> "Hard rules".
# =============================================================================

"""Generate missing sample DOCX files for AI Recruitment System."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

SAMPLES_DIR = r"C:\AI-Recruitment-System\samples"
os.makedirs(SAMPLES_DIR, exist_ok=True)


def create_interview_questions():
    doc = Document()
    title = doc.add_heading("Interview Guide — Senior Software Engineer", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Candidate: Rahul Sharma | Requisition: REQ-2026-SE-0142 | "
        "Panel: Engineering + HR"
    )
    doc.add_paragraph(
        "Generated: 2026-08-03 | Version: 1.0 | Match Score: 87/100"
    )

    doc.add_heading("Role Context", level=1)
    doc.add_paragraph(
        "Senior Software Engineer on the Platform Engineering team. "
        "Focus: distributed systems, Python/Go, AWS, mentoring junior engineers."
    )

    doc.add_heading("Technical Questions (45 min)", level=1)
    questions = [
        (
            "System Design",
            "Design a resume ingestion pipeline that handles 10,000 applications/day. "
            "Discuss queueing, idempotency, and failure recovery.",
        ),
        (
            "Python",
            "Explain how you would optimize a slow pandas DataFrame operation "
            "processing 2M rows. What profiling tools would you use?",
        ),
        (
            "AWS",
            "Walk through how you migrated a monolithic API to ECS Fargate. "
            "What trade-offs did you consider?",
        ),
        (
            "Code Review",
            "Given a PR that adds caching to an API endpoint, "
            "what would you look for beyond correctness?",
        ),
    ]
    for i, (cat, q) in enumerate(questions, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. [{cat}] ").bold = True
        p.add_run(q)

    doc.add_heading("Behavioral Questions (30 min)", level=1)
    behavioral = [
        "Tell me about a time you disagreed with a technical decision. "
        "How did you handle it?",
        "Describe a situation where you mentored a junior engineer "
        "through a production incident.",
        "Give an example of delivering under a tight deadline "
        "without compromising quality.",
    ]
    for i, q in enumerate(behavioral, 1):
        doc.add_paragraph(f"{i}. {q}")

    doc.add_heading("Evaluation Rubric", level=1)
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    for j, h in enumerate(["Competency", "Weight", "Score (1-5)", "Notes"]):
        table.rows[0].cells[j].text = h
    rows = [
        ("Technical Depth", "30%", "", ""),
        ("System Design", "25%", "", ""),
        ("Communication", "20%", "", ""),
        ("Culture Fit", "25%", "", ""),
    ]
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val

    path = os.path.join(SAMPLES_DIR, "Interview_Questions.docx")
    doc.save(path)
    print(f"Created {path}")


def create_interview_feedback():
    doc = Document()
    title = doc.add_heading("Interview Feedback Form", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Candidate: Rahul Sharma | Role: Senior Software Engineer | "
        "Interviewer: Priya Mehta (Engineering Manager)"
    )
    doc.add_paragraph(
        "Interview Date: 2026-08-10 | Duration: 75 minutes | "
        "Recommendation: STRONG HIRE"
    )

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        "Rahul demonstrated strong technical depth in distributed systems and Python. "
        "His AWS migration story was detailed and showed good architectural judgment. "
        "Communication was clear and structured. Minor gap in Go experience but showed "
        "willingness to ramp up."
    )

    doc.add_heading("Competency Scores", level=1)
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    for j, h in enumerate(["Competency", "Score (1-5)", "Evidence"]):
        table.rows[0].cells[j].text = h
    scores = [
        ("Technical Depth", "4.5", "Strong Python, solid AWS. Explained caching trade-offs well."),
        ("System Design", "4.0", "Good pipeline design. Could improve on cost optimization discussion."),
        ("Communication", "4.5", "Structured answers, good use of examples."),
        ("Culture Fit", "4.0", "Collaborative mindset. Values mentoring."),
    ]
    for i, row in enumerate(scores, 1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val

    doc.add_heading("Strengths", level=1)
    for s in [
        "Production incident handling at scale",
        "Clear technical communication",
        "Mentoring experience with 3 junior engineers",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("Areas for Development", level=1)
    for a in [
        "Go language proficiency (currently learning)",
        "Cost optimization in cloud architecture",
    ]:
        doc.add_paragraph(a, style="List Bullet")

    doc.add_heading("Final Recommendation", level=1)
    p = doc.add_paragraph()
    p.add_run("STRONG HIRE").bold = True
    doc.add_paragraph("Confidence: High | Comp band alignment: L5 Senior Engineer")

    path = os.path.join(SAMPLES_DIR, "Interview_Feedback.docx")
    doc.save(path)
    print(f"Created {path}")


if __name__ == "__main__":
    create_interview_questions()
    create_interview_feedback()
