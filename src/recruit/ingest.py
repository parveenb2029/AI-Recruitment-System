"""Document ingest — the first stage of WF-03.

Turns a file on disk into a `Document`: validated, text-extracted, hashed.
Everything downstream consumes `Document.text`, so this module is where messy
real-world files stop being messy.

Implements the file-validation rules in `03_Extracted_Data/Validation.md` §2 and
the reason codes in `Error_Handling.md` §2.

    python -m recruit.ingest samples/Rahul_Sharma_Resume.pdf
"""

from __future__ import annotations

import argparse
import hashlib
import json
import sys
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

from .errors import (
    CorruptDocument,
    EmptySource,
    EncryptedDocument,
    FileTooLarge,
    InvalidDocument,
    OCRUnavailable,
    RecruitError,
    UnsupportedFormat,
)

# Magic bytes per extension. Validation.md §2 requires the MIME type to match the
# extension — a .pdf that is secretly a .exe must not reach the parser.
MAGIC = {
    ".pdf": [b"%PDF-"],
    ".docx": [b"PK\x03\x04"],          # zip container
    ".doc": [b"\xd0\xcf\x11\xe0"],     # OLE2 compound file
    ".txt": [],                        # no signature; validated by decoding
}

DEFAULT_MAX_MB = 10
DEFAULT_MIN_CHARS = 100


@dataclass
class Document:
    """An ingested source document, ready for extraction."""

    path: Path
    text: str
    content_sha256: str
    extension: str
    size_bytes: int
    pages: int = 0
    ocr_used: bool = False
    # Validation.md §2 requires a virus scan. No scanner is wired up yet, so this
    # records the truth rather than implying a check that did not happen.
    # See CLAUDE.md, Deferred work.
    virus_scanned: bool = False
    warnings: list[str] = field(default_factory=list)

    @property
    def char_count(self) -> int:
        return len(self.text)

    def summary(self) -> dict[str, object]:
        return {
            "path": str(self.path),
            "extension": self.extension,
            "size_bytes": self.size_bytes,
            "pages": self.pages,
            "char_count": self.char_count,
            "ocr_used": self.ocr_used,
            "virus_scanned": self.virus_scanned,
            "content_sha256": self.content_sha256,
            "warnings": self.warnings,
        }


# -- validation ---------------------------------------------------------------
def _validate_file(path: Path, max_mb: int, accepted: list[str]) -> tuple[bytes, str]:
    if not path.is_file():
        raise InvalidDocument(f"No such file: {path}")

    extension = path.suffix.lower()
    if extension not in accepted:
        raise UnsupportedFormat(
            f"Extension {extension or '(none)'} is not accepted.",
            detail=f"Accepted: {', '.join(accepted)}",
        )

    size = path.stat().st_size
    if size == 0:
        raise CorruptDocument("File is empty (0 bytes).")
    if size > max_mb * 1024 * 1024:
        raise FileTooLarge(
            f"File is {size / 1024 / 1024:.1f} MB, limit is {max_mb} MB."
        )

    raw = path.read_bytes()

    signatures = MAGIC.get(extension, [])
    if signatures and not any(raw.startswith(sig) for sig in signatures):
        raise InvalidDocument(
            f"File content does not match its {extension} extension.",
            detail=f"Starts with {raw[:8]!r}; expected one of "
                   f"{[s[:8] for s in signatures]}",
        )

    return raw, extension


# -- extraction ---------------------------------------------------------------
def _extract_pdf(path: Path, raw: bytes) -> tuple[str, int]:
    try:
        from pypdf import PdfReader
        from pypdf.errors import PdfReadError
    except ImportError:  # pragma: no cover
        raise RecruitError(
            "pypdf is not installed.", detail="pip install pypdf"
        ) from None

    try:
        reader = PdfReader(str(path))
    except PdfReadError as exc:
        raise CorruptDocument("PDF could not be parsed.", detail=str(exc)) from exc
    except Exception as exc:  # noqa: BLE001 - pypdf raises assorted types
        raise CorruptDocument("PDF could not be opened.", detail=str(exc)) from exc

    if reader.is_encrypted:
        # An empty user password is common and harmless; try it before giving up.
        try:
            if reader.decrypt("") == 0:
                raise EncryptedDocument("PDF is password-protected.")
        except EncryptedDocument:
            raise
        except Exception as exc:  # noqa: BLE001
            raise EncryptedDocument(
                "PDF is encrypted and could not be opened.", detail=str(exc)
            ) from exc

    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 - one bad page must not lose the rest
            parts.append("")
    return "\n".join(parts).strip(), len(reader.pages)


def _extract_docx(path: Path) -> tuple[str, int]:
    try:
        import docx
    except ImportError:  # pragma: no cover
        raise RecruitError(
            "python-docx is not installed.", detail="pip install python-docx"
        ) from None

    try:
        document = docx.Document(str(path))
    except zipfile.BadZipFile as exc:
        raise CorruptDocument("DOCX container is not a valid zip.", detail=str(exc)) from exc
    except Exception as exc:  # noqa: BLE001
        raise CorruptDocument("DOCX could not be opened.", detail=str(exc)) from exc

    parts = [p.text for p in document.paragraphs]
    # Tables carry real content in resumes — skills grids, dates. Do not drop them.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts).strip(), 0


def _extract_txt(raw: bytes) -> tuple[str, int]:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return raw.decode(encoding).strip(), 0
        except UnicodeDecodeError:
            continue
    raise InvalidDocument("Text file could not be decoded in any known encoding.")


def _ocr_pdf(path: Path) -> str:
    """OCR fallback. Requires Tesseract, which is a system binary, not a wheel."""
    try:
        import pytesseract
        from pdf2image import convert_from_path
    except ImportError as exc:
        raise OCRUnavailable(
            "OCR is required for this document but no OCR engine is installed.",
            detail=f"{exc}. Install Tesseract plus: pip install pytesseract pdf2image",
        ) from exc

    try:
        images = convert_from_path(str(path))
    except Exception as exc:  # noqa: BLE001
        raise OCRUnavailable(
            "Could not rasterize the PDF for OCR.",
            detail=f"{exc}. Poppler may be missing.",
        ) from exc

    return "\n".join(pytesseract.image_to_string(image) for image in images).strip()


# -- public API ---------------------------------------------------------------
def load(
    path: str | Path,
    *,
    max_mb: int = DEFAULT_MAX_MB,
    min_chars: int = DEFAULT_MIN_CHARS,
    accepted: list[str] | None = None,
    allow_ocr: bool = True,
) -> Document:
    """Validate, extract, and hash a source document.

    Raises a typed `RecruitError` subclass on every failure — never a bare
    traceback. Callers branch on `error.code`.
    """
    path = Path(path)
    accepted = accepted or list(MAGIC.keys())

    raw, extension = _validate_file(path, max_mb, accepted)

    # Hash the SOURCE BYTES, not the extracted text. Extraction is a function of
    # the library version; the bytes are not. This is the idempotency key
    # (Execution_Flow.md §8), so it must be stable across pypdf upgrades.
    content_sha256 = hashlib.sha256(raw).hexdigest()

    warnings: list[str] = []
    ocr_used = False

    if extension == ".pdf":
        text, pages = _extract_pdf(path, raw)
    elif extension == ".docx":
        text, pages = _extract_docx(path)
    elif extension == ".txt":
        text, pages = _extract_txt(raw)
    else:  # .doc — legacy binary format
        raise UnsupportedFormat(
            "Legacy .doc files are not supported.",
            detail="Save as .docx or .pdf and re-upload.",
        )

    # Validation.md §2: below the threshold, take the OCR path.
    if len(text) < min_chars:
        if extension == ".pdf" and allow_ocr:
            warnings.append(
                f"Text layer yielded {len(text)} chars (< {min_chars}); used OCR."
            )
            text = _ocr_pdf(path)
            ocr_used = True
        if len(text) < min_chars:
            raise EmptySource(
                f"Only {len(text)} extractable characters (minimum {min_chars}).",
                detail="Document may be a scan, an image, or effectively blank.",
            )

    return Document(
        path=path,
        text=text,
        content_sha256=content_sha256,
        extension=extension,
        size_bytes=len(raw),
        pages=pages,
        ocr_used=ocr_used,
        virus_scanned=False,
        warnings=warnings,
    )


# -- CLI ----------------------------------------------------------------------
def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        prog="python -m recruit.ingest", description=__doc__
    )
    parser.add_argument("file", type=Path)
    parser.add_argument("--json", action="store_true", help="Emit summary as JSON.")
    parser.add_argument("--text", action="store_true", help="Print all text, not a preview.")
    parser.add_argument("--no-ocr", action="store_true", help="Disable OCR fallback.")
    parser.add_argument("--max-mb", type=int, default=DEFAULT_MAX_MB)
    args = parser.parse_args(argv)

    try:
        document = load(args.file, max_mb=args.max_mb, allow_ocr=not args.no_ocr)
    except RecruitError as error:
        if args.json:
            print(json.dumps(error.as_dict(), indent=2))
        else:
            print(str(error), file=sys.stderr)
            print(f"  recovery: {error.recovery}", file=sys.stderr)
        return 1

    if args.json:
        print(json.dumps(document.summary(), indent=2))
        return 0

    summary = document.summary()
    print(f"  file          {summary['path']}")
    print(f"  format        {summary['extension']}  ({summary['size_bytes']:,} bytes)")
    if summary["pages"]:
        print(f"  pages         {summary['pages']}")
    print(f"  characters    {summary['char_count']:,}")
    print(f"  ocr used      {summary['ocr_used']}")
    print(f"  virus scanned {summary['virus_scanned']}   (no scanner wired up yet)")
    print(f"  sha256        {summary['content_sha256']}")
    for warning in document.warnings:
        print(f"  warning       {warning}")

    print("\n  --- text " + "-" * 60)
    body = document.text if args.text else document.text[:600]
    for line in body.splitlines():
        print(f"  {line}")
    if not args.text and len(document.text) > 600:
        print(f"  ... [{len(document.text) - 600:,} more characters; use --text]")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
